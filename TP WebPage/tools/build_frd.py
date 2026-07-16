from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "frd"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "Trainee_Portal_Functional_Requirements_Document_v1.0.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
PALE = "F4F6F9"
GRID = "B8C4CE"
MUTED = "5B6573"
GREEN = "2E7D32"
AMBER = "9A6700"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def border_paragraph(paragraph, color=GRID, fill=PALE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def add_field_toc_note(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    border_paragraph(p, color="C6D4E1", fill=LIGHT_BLUE)
    r = p.add_run("Document navigation note: ")
    set_run_font(r, size=8.5, bold=True, color=NAVY)
    r = p.add_run("The table of contents is static for reliable PDF export. Page numbers match the verified final rendering.")
    set_run_font(r, size=8.5, color=MUTED)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string("20252B")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 10),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 17, BLUE, 0, 9),
        ("Heading 2", 12.5, NAVY, 9, 5),
        ("Heading 3", 10.5, NAVY, 6, 3),
    ):
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = style_name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    # Running furniture. The first page is intentionally clear.
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("TRAINEE PORTAL  |  FUNCTIONAL REQUIREMENTS DOCUMENT")
    set_run_font(r, size=8, bold=True, color=MUTED)
    fp = section.footer.paragraphs[0]
    add_page_number(fp)
    return doc


def page_break(doc):
    doc.add_page_break()


def add_page_title(doc, number, title, subtitle=None):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(f"{number}. {title}")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(subtitle)
        set_run_font(r2, size=9, italic=True, color=MUTED)


def add_text(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items, compact=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2 if compact else 4)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(item)
        set_run_font(r, size=9.5 if compact else 10)


def add_definition_rows(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Definition / Value"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9, bold=True, color=WHITE)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    set_run_font(run, size=9, bold=(idx == 0), color=NAVY if idx == 0 else None)
    set_table_geometry(table, [2200, 7160])
    return table


def add_requirements(doc, requirements, source_note=None):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["ID", "Functional requirement", "User(s)", "Implementation evidence"]
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = label
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8.3, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for req_id, req, users, evidence in requirements:
        cells = table.add_row().cells
        values = [req_id, req, users, evidence]
        for idx, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            if idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=7.8 if idx in (0, 2, 3) else 8.2, bold=(idx == 0), color=NAVY if idx == 0 else None)
    set_table_geometry(table, [780, 5330, 1150, 2100])
    if source_note:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Source basis: ")
        set_run_font(r, size=7.6, bold=True, color=MUTED)
        r = p.add_run(source_note)
        set_run_font(r, size=7.6, color=MUTED)
    return table


def add_rules(doc, rules, title="Business rules"):
    doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, label in enumerate(("Rule", "Statement", "Source / rationale")):
        cell = table.rows[0].cells[i]
        cell.text = label
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for rid, statement, source in rules:
        cells = table.add_row().cells
        for idx, value in enumerate((rid, statement, source)):
            cells[idx].text = value
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT_BLUE)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=8.1, bold=(idx == 0), color=NAVY if idx == 0 else None)
    set_table_geometry(table, [900, 5780, 2680])


def add_screenshot_placeholder(doc, label, suggested_capture):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    border_paragraph(p, color="AAB7C4", fill="F7F9FB")
    r = p.add_run(f"SCREENSHOT PLACEHOLDER - {label}\n")
    set_run_font(r, size=9, bold=True, color=BLUE)
    r = p.add_run(f"Suggested capture: {suggested_capture}\nInsert final application screenshot here; retain sensitive/demo data only.")
    set_run_font(r, size=8.3, color=MUTED)


def add_gap_callout(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    border_paragraph(p, color="D5A021", fill="FFF8E6")
    r = p.add_run(f"{title}: ")
    set_run_font(r, size=9, bold=True, color=AMBER)
    r = p.add_run(text)
    set_run_font(r, size=9, color="5A4300")


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FUNCTIONAL REQUIREMENTS\nDOCUMENT")
    set_run_font(r, size=29, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Trainee Portal")
    set_run_font(r, size=20, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("As-Is Functional Specification Based on Implemented Source Code")
    set_run_font(r, size=11, italic=True, color=MUTED)
    p.paragraph_format.space_after = Pt(62)
    add_definition_rows(doc, [
        ("Document version", "1.0"),
        ("Prepared as", "Business Analyst deliverable - internship/final-year project"),
        ("Prepared date", "13 July 2026"),
        ("Project", "Trainee Portal Web Application"),
        ("Confidentiality", "Academic / internal project use"),
    ])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    r = p.add_run("Prepared from the repository implementation; no unverified features have been added.")
    set_run_font(r, size=9, color=MUTED)


def add_control_page(doc):
    add_page_title(doc, "", "Document Control".strip())
    doc.paragraphs[-1].text = "Document Control"
    doc.paragraphs[-1].style = "Heading 1"
    doc.add_paragraph("Version history", style="Heading 2")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(("Version", "Date", "Author / Role", "Change summary")):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], NAVY)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_run_font(run, size=9, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("0.1", "13 Jul 2026", "Business Analyst", "Initial implementation review and requirement inventory."),
        ("1.0", "13 Jul 2026", "Business Analyst", "Final FRD with verified modules, gaps, business rules and NFRs."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs:
                set_run_font(run, size=8.7)
    set_table_geometry(table, [1000, 1500, 1900, 4960])
    doc.add_paragraph("Review and approval", style="Heading 2")
    add_definition_rows(doc, [
        ("Prepared by", "____________________________  Date: __________"),
        ("Reviewed by", "____________________________  Date: __________"),
        ("Approved by", "____________________________  Date: __________"),
    ])
    doc.add_paragraph("Document conventions", style="Heading 2")
    add_bullets(doc, [
        "Implemented - supported by frontend/backend code and, where applicable, automated tests.",
        "Frontend only - visible behavior exists, but there is no persistent backend endpoint.",
        "Configuration only - the screen stores a setting locally but no scheduler/job executes it.",
        "Known gap - explicitly absent or incomplete in the current repository.",
    ], compact=True)
    add_gap_callout(doc, "Scope caution", "This FRD specifies the current codebase. It does not certify a live production deployment or a real PostgreSQL migration run.")


TOC_ENTRIES = [
    ("1", "Introduction", 4), ("2", "Project Overview", 5), ("3", "User Roles and Access", 6),
    ("4", "Authentication and Account Access", 7), ("5", "User and Profile Management", 8),
    ("6", "Batch Management", 9), ("7", "Assignment Management", 10),
    ("8", "Submission and Grading", 11), ("9", "Sessions and Calendar", 12),
    ("10", "Attendance Management", 13), ("11", "Resource Library", 14),
    ("12", "Feedback Management", 15), ("13", "Notifications and Audit Logs", 16),
    ("14", "Dashboards and Analytics", 17), ("15", "Announcements and Reports", 18),
    ("16", "Common User Experience Functions", 19), ("17", "Business Rules - Access and Identity", 20),
    ("18", "Business Rules - Training Operations", 21), ("19", "Non-Functional Requirements - Security", 22),
    ("20", "Non-Functional Requirements - Performance and Reliability", 23),
    ("21", "Non-Functional Requirements - Usability and Maintainability", 24),
    ("22", "Data, Integrations and Constraints", 25), ("23", "Known Gaps and Future Enhancements", 26),
    ("24", "Conclusion and Verification Summary", 27),
]


def add_toc(doc):
    doc.add_paragraph("Table of Contents", style="Heading 1")
    for num, title, page in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(f"{num}. {title}\t{page}")
        set_run_font(r, size=9.2, bold=(num in {"1", "4", "17", "19", "23", "24"}), color=NAVY if num in {"1", "4", "17", "19", "23", "24"} else None)
    add_field_toc_note(doc)


def build():
    doc = setup_document()
    add_cover(doc)
    page_break(doc)
    add_control_page(doc)
    page_break(doc)
    add_toc(doc)

    page_break(doc)
    add_page_title(doc, "1", "Introduction")
    add_text(doc, "Purpose: This Functional Requirements Document (FRD) defines what the Trainee Portal currently does, who can use each capability, and the principal rules that govern the implemented workflows.", "Purpose:")
    add_text(doc, "The document is intentionally concise and suitable for an internship or final-year project. It is an as-is specification derived from the repository rather than a wish list.")
    doc.add_paragraph("Objectives", style="Heading 2")
    add_bullets(doc, [
        "Provide role-specific workspaces for administrators, facilitators and trainees.",
        "Centralize batches, assignments, submissions, sessions, attendance, learning resources and feedback.",
        "Provide operational visibility through metrics, notifications, audit activity and downloadable reports.",
        "Apply authenticated and role-aware access to sensitive actions and files.",
    ])
    doc.add_paragraph("Scope and method", style="Heading 2")
    add_bullets(doc, [
        "In scope: React/Vite frontend, Express/TypeScript API, Prisma/PostgreSQL data model, storage abstraction, validation, tests and deployment configuration.",
        "Out of scope: proposed features without code, third-party vendor commitments, production service-level agreements and business processes not represented in the repository.",
        "Evidence date: repository reviewed on 13 July 2026. Production build succeeded; backend automated suite passed 92 tests in 20 files.",
    ])
    add_gap_callout(doc, "Interpretation", "Where a screen exists without persistence or a backend route, the requirement is labeled Frontend only or Configuration only instead of being described as fully implemented.")

    page_break(doc)
    add_page_title(doc, "2", "Project Overview")
    add_text(doc, "The Trainee Portal is a browser-based training administration system implemented as a TypeScript monorepo. It combines role-specific dashboards with a REST API and relational data model.")
    add_definition_rows(doc, [
        ("Frontend", "React 18, TypeScript, Vite, React Router, Zustand and Tailwind CSS"),
        ("Backend", "Node.js, Express and TypeScript REST API"),
        ("Persistence", "PostgreSQL accessed through Prisma ORM"),
        ("Files", "Pluggable local or Amazon S3 storage provider"),
        ("Authentication", "JWT access token plus rotating refresh-token cookie"),
        ("Deployment", "Vercel frontend; Railway or Render backend; Neon-ready database configuration"),
    ])
    doc.add_paragraph("Logical flow", style="Heading 2")
    add_bullets(doc, [
        "A user signs in or accepts an administrator-issued invitation.",
        "The application directs the authenticated user to the dashboard for their assigned role.",
        "The frontend requests protected REST resources; the API validates identity, role, input and ownership.",
        "Business records are stored in PostgreSQL and uploaded files are stored through the configured provider.",
        "Selected changes create audit entries that also act as the source for user notifications.",
    ], compact=True)
    add_screenshot_placeholder(doc, "System overview", "application login page beside the three role dashboards, or a simple deployment architecture diagram")

    page_break(doc)
    add_page_title(doc, "3", "User Roles and Access")
    add_definition_rows(doc, [
        ("Admin", "System-wide administrator. Manages users, invitations, batches, global assignments, sessions, resources, feedback review, reports and audit visibility."),
        ("Facilitator", "Training operator. Manages assigned batches, creates assignments/sessions/resources, reviews submissions, marks attendance and gives trainee feedback."),
        ("Trainee", "Learner. Views enrolled batches and assigned work, submits files, joins sessions, uses verified resources, reviews grades and provides facilitator feedback."),
    ])
    doc.add_paragraph("Role boundaries", style="Heading 2")
    add_bullets(doc, [
        "The frontend protects Admin, Facilitator and Trainee routes through RequireAuth role checks.",
        "The backend separately protects privileged routes with authentication and role middleware; ownership and enrollment rules are enforced inside services.",
        "Admin access is generally global. Facilitator actions are normally limited to owned/assigned objects. Trainee access is normally limited to active enrollment or ownership.",
        "The schema contains permissions and role-permission mappings; current route authorization primarily uses the three role names.",
    ])
    add_gap_callout(doc, "Access caveat", "The assignment-submission roster endpoint is intentionally readable by any authenticated user in the current service implementation. This is broader than a least-privilege academic portal would normally require.")
    add_screenshot_placeholder(doc, "Role dashboards", "one screenshot per role showing the left navigation and dashboard header")

    page_break(doc)
    add_page_title(doc, "4", "Authentication and Account Access")
    add_requirements(doc, [
        ("FR-001", "The system shall authenticate an active user by email and password and return the user profile with an access token.", "All", "Implemented - API + UI"),
        ("FR-002", "The system shall support a Remember Me choice when creating the refresh-token session.", "All", "Implemented - API + UI"),
        ("FR-003", "The system shall rotate a valid refresh token and issue a new access/refresh token pair.", "All", "Implemented - tested"),
        ("FR-004", "The system shall revoke the current refresh token and clear the refresh cookie on logout.", "All", "Implemented - API + UI"),
        ("FR-005", "The system shall return the current authenticated user and restore the session during application bootstrap.", "All", "Implemented - API + UI"),
        ("FR-006", "The system shall allow an administrator to create an invitation for an Admin, Facilitator or Trainee email address.", "Admin", "Implemented; email gap"),
        ("FR-007", "The system shall allow a recipient with a valid, pending and unexpired invite token to create an account password and activate the account.", "Invitee", "Implemented - API + UI"),
        ("FR-008", "The system shall allow a signed-in user to change password after validating the current password.", "All", "Implemented - API + UI"),
        ("FR-009", "The login screen shall expose demo entry points for the three roles using in-memory/session data.", "Demo user", "Implemented - demo only"),
    ], "auth.routes.ts, auth.service.ts, auth.validator.ts, authStore.ts, LoginPage.tsx, InvitePage.tsx, token-refresh and auth tests.")
    add_gap_callout(doc, "Known security gap", "Forgot Password accepts email plus a new password without a verification token. It exists in code but should not be treated as a production-ready recovery workflow.")
    add_screenshot_placeholder(doc, "Authentication", "login screen and invite acceptance screen")

    page_break(doc)
    add_page_title(doc, "5", "User and Profile Management")
    add_requirements(doc, [
        ("FR-010", "The system shall allow an administrator to list users with role, active-status, search, sorting and pagination filters.", "Admin", "Implemented - API"),
        ("FR-011", "The system shall allow an authenticated user to view their own profile.", "All", "Implemented - API + UI"),
        ("FR-012", "The system shall allow a user to update their name, email, phone and location.", "All", "Implemented - API + UI"),
        ("FR-013", "The system shall allow an administrator to change a user's role, active state and company/profile metadata.", "Admin", "Implemented - API"),
        ("FR-014", "The system shall allow an administrator to soft-delete a user and revoke outstanding refresh tokens.", "Admin", "Implemented - API"),
        ("FR-015", "The account settings page shall display company/batch metadata as read-only contextual information by role.", "All", "Implemented - UI"),
    ], "users.routes.ts, users.service.ts, UserProfile schema, AccountSettingsPage.tsx and user service/store.")
    doc.add_paragraph("Current limitations", style="Heading 2")
    add_bullets(doc, [
        "No self-registration route is implemented; account creation is invitation-based or seeded for demo use.",
        "Avatar metadata exists in the data model, but no complete avatar upload endpoint is exposed in the reviewed routes.",
        "No user-facing administration screen was found for editing the role-permission matrix.",
    ])
    add_screenshot_placeholder(doc, "Account settings", "profile information form and role-specific details panel")

    page_break(doc)
    add_page_title(doc, "6", "Batch Management")
    add_requirements(doc, [
        ("FR-016", "The system shall list non-deleted batches with search, filters, sorting and pagination.", "All", "Implemented - API + UI"),
        ("FR-017", "The system shall scope batch lists to a specified facilitator or trainee when the relevant filter is supplied.", "All", "Implemented - tested"),
        ("FR-018", "The system shall allow an administrator to create a batch with a unique code, name, program, track, dates, status and optional facilitator.", "Admin", "Implemented - tested"),
        ("FR-019", "The system shall allow an administrator to update or soft-delete a batch.", "Admin", "Implemented - tested"),
        ("FR-020", "The system shall calculate batch metrics for trainee count, average grade, completion, submission, attendance and feedback rating.", "Admin, Facilitator", "Implemented - API"),
        ("FR-021", "The system shall allow Admin or Facilitator users to list trainees in an accessible batch.", "Admin, Facilitator", "Implemented - scoped"),
        ("FR-022", "The system shall allow Admin or Facilitator users to enroll or remove a trainee from a batch without deleting historical enrollment rows.", "Admin, Facilitator", "Implemented - API"),
        ("FR-023", "The system shall provide per-trainee batch statistics including assignment progress, average grade, attendance, latest submission and feedback status.", "Admin, owner Facilitator", "Implemented - tested"),
    ], "batches.routes.ts, batches.service.ts, Batch and BatchTrainee schema, dashboards and batch scope/CRUD tests.")
    add_screenshot_placeholder(doc, "Batch management", "Admin batch table and Facilitator batch-detail performance page")

    page_break(doc)
    add_page_title(doc, "7", "Assignment Management")
    add_requirements(doc, [
        ("FR-024", "The system shall list active assignments with batch/status filters, search, sorting and pagination.", "All", "Implemented - API + UI"),
        ("FR-025", "The system shall allow Admin or Facilitator users to create an assignment for one or more batches.", "Admin, Facilitator", "Implemented - tested"),
        ("FR-026", "An assignment shall contain a title, optional description, deadline, status and optional instruction file.", "Admin, Facilitator", "Implemented - validated"),
        ("FR-027", "The creator shall be recorded as the assignment facilitator, and the first selected batch shall remain the primary batch for compatibility.", "System", "Implemented - service/schema"),
        ("FR-028", "The system shall allow an administrator or owning facilitator to update assignment data, batch links or instruction file.", "Admin, owner Facilitator", "Implemented - authorized"),
        ("FR-029", "The system shall allow an administrator or owning facilitator to soft-delete an assignment.", "Admin, owner Facilitator", "Implemented - authorized"),
        ("FR-030", "The system shall allow an administrator, owning/co-facilitator or enrolled trainee to view the instruction attachment.", "Scoped users", "Implemented - tested"),
        ("FR-031", "The assignment dashboard shall support selection of multiple assignments and bulk deadline extension in the client workflow.", "Admin, Facilitator", "Implemented - UI/API calls"),
    ], "assignments.routes.ts, assignments.service.ts, assignment validators/schema, Admin/Facilitator dashboards and multi-batch tests.")
    add_screenshot_placeholder(doc, "Assignment management", "assignment table, multi-batch selection and create-assignment modal")

    page_break(doc)
    add_page_title(doc, "8", "Submission and Grading")
    add_requirements(doc, [
        ("FR-032", "The system shall display one roster row for every enrolled trainee, including a Not Started placeholder when no submission exists.", "Authenticated", "Implemented - API"),
        ("FR-033", "The system shall allow an actively enrolled trainee to submit their own work for an assigned assignment.", "Trainee", "Implemented - authorized"),
        ("FR-034", "The system shall mark a new submission Under Review or Late according to the assignment deadline.", "System", "Implemented - service"),
        ("FR-035", "The system shall upsert one submission per trainee and assignment when the trainee submits again.", "Trainee", "Implemented - schema/service"),
        ("FR-036", "The system shall allow the trainee to attach a file to their own submission.", "Trainee", "Implemented - upload API"),
        ("FR-037", "The system shall allow replacement of a current attachment before the deadline and retain the superseded attachment history.", "Trainee", "Implemented - tested"),
        ("FR-038", "The system shall reject replacement of an existing attachment after the deadline, while allowing a first late attachment.", "Trainee", "Implemented - tested"),
        ("FR-039", "The system shall allow an administrator or owning facilitator to enter a grade from 0 to 100 and optional written feedback.", "Admin, owner Facilitator", "Implemented - API + UI"),
        ("FR-040", "Submission attachments shall be downloadable only by the trainee owner, administrator or owning facilitator.", "Scoped users", "Implemented - tested"),
    ], "submissions routes/service/validator/schema, AssignmentDetailPage.tsx and submission authorization/resubmission tests.")
    add_gap_callout(doc, "UI placeholder", "The Facilitator dashboard includes a text-only 'submission summary' download helper labeled as a placeholder; actual uploaded submission files are handled through the secured attachment route.")

    page_break(doc)
    add_page_title(doc, "9", "Sessions and Calendar")
    add_requirements(doc, [
        ("FR-041", "The system shall list sessions with batch/status filters, search, sorting and pagination.", "All", "Implemented - API + UI"),
        ("FR-042", "The system shall allow Admin or Facilitator users to schedule a batch session with title, date/time, platform, optional meeting URL and status.", "Admin, Facilitator", "Implemented - API + UI"),
        ("FR-043", "The system shall allow an administrator or owning facilitator to update, reschedule, cancel or soft-delete a session.", "Admin, owner Facilitator", "Implemented - API + UI"),
        ("FR-044", "The system shall support Google Meet, Microsoft Teams, Zoom and Other as session platforms.", "Admin, Facilitator", "Implemented - enum/validation"),
        ("FR-045", "The system shall expose session states Upcoming, Live, Completed, Cancelled and Rescheduled.", "All", "Implemented - enum/validation"),
        ("FR-046", "The system shall provide a combined calendar feed of sessions and assignment deadlines.", "All", "Implemented - tested"),
        ("FR-047", "Calendar results shall be scoped to all batches for Admin, managed batches for Facilitator and active enrollments for Trainee.", "All", "Implemented - tested"),
        ("FR-048", "A trainee shall be able to open a valid meeting link from an eligible session.", "Trainee", "Implemented - UI"),
    ], "sessions/calendar routes and services, schemas, SessionsCalendarView.tsx and calendar tests.")
    add_screenshot_placeholder(doc, "Sessions and calendar", "list/calendar toggle with an upcoming session and assignment deadline")

    page_break(doc)
    add_page_title(doc, "10", "Attendance Management")
    add_requirements(doc, [
        ("FR-049", "The system shall allow Admin or Facilitator users to retrieve attendance for a session.", "Admin, Facilitator", "Implemented - API"),
        ("FR-050", "The system shall allow Admin or Facilitator users to bulk mark attendance for one or more trainees in a session.", "Admin, Facilitator", "Implemented - API"),
        ("FR-051", "The system shall allow an administrator or owning facilitator to correct a single attendance record.", "Admin, owner Facilitator", "Implemented - API"),
        ("FR-052", "Attendance status shall be Present, Absent, Late or Excused.", "System", "Implemented - enum/validation"),
        ("FR-053", "Attendance shall be unique for each session and trainee combination and record who marked it.", "System", "Implemented - schema"),
        ("FR-054", "Batch analytics shall treat Present and Late records as attended when calculating attendance rate.", "System", "Implemented - service/test"),
    ], "attendance routes/service/validator, Attendance schema, batch metrics and tests.")
    add_gap_callout(doc, "Demo limitation", "Demo Mode returns an empty attendance collection because no per-trainee attendance fixture is included. The production API and data model are implemented.")
    add_screenshot_placeholder(doc, "Attendance", "session attendance roster with status controls and summary counts")

    page_break(doc)
    add_page_title(doc, "11", "Resource Library")
    add_requirements(doc, [
        ("FR-055", "The system shall list resources with batch, category, verification, search, sorting and pagination filters.", "All", "Implemented - API + UI"),
        ("FR-056", "The system shall allow Admin or Facilitator users to upload a file with title, category, version and optional batch scope.", "Admin, Facilitator", "Implemented - upload API"),
        ("FR-057", "The system shall record file storage key, MIME type, byte size, uploader and download count.", "System", "Implemented - schema/service"),
        ("FR-058", "The uploader or an administrator shall be able to update resource metadata or soft-delete a resource.", "Admin, owner Facilitator", "Implemented - authorized"),
        ("FR-059", "An administrator shall be able to set the verification status used by the resource repository UI.", "Admin", "Implemented - update API/UI"),
        ("FR-060", "The system shall stream an available resource file and increment its download count.", "All", "Implemented - API"),
        ("FR-061", "The Admin interface shall support bulk verification and bulk deletion through individual secured API calls.", "Admin", "Implemented - UI orchestration"),
    ], "resources routes/service/validator/schema, storage providers and Admin/Facilitator/Trainee dashboards.")
    add_gap_callout(doc, "Visibility rule caveat", "The UI text says resources become visible to trainees once verified, but the generic list API does not automatically force verified=true by caller role. Enforcement depends on the frontend query and should be hardened server-side.")
    add_screenshot_placeholder(doc, "Resource library", "resource filters, verified badge, upload dialog and download action")

    page_break(doc)
    add_page_title(doc, "12", "Feedback Management")
    add_requirements(doc, [
        ("FR-062", "The system shall list feedback with batch, trainee, facilitator, direction and rating/date sorting filters.", "Scoped users", "Implemented - API + UI"),
        ("FR-063", "The system shall allow Admin or Facilitator users to submit rating-based feedback about an actively enrolled trainee.", "Admin, Facilitator", "Implemented - tested"),
        ("FR-064", "The system shall allow a trainee to submit feedback only about the facilitator assigned to an actively enrolled batch.", "Trainee", "Implemented - tested"),
        ("FR-065", "Feedback shall include a category, integer rating from 1 to 5 and optional comment.", "All", "Implemented - validated"),
        ("FR-066", "The system shall stamp feedback direction automatically as Facilitator-to-Trainee or Trainee-to-Facilitator.", "System", "Implemented - tested"),
        ("FR-067", "Feedback records shall be append-only; no update or delete endpoint shall be exposed.", "All", "Implemented - route design"),
        ("FR-068", "A trainee's feedback list shall be restricted to that trainee even if another trainee filter is requested.", "Trainee", "Implemented - tested"),
    ], "feedback routes/service/validator/schema, feedback dashboards and feedback-direction tests.")
    add_screenshot_placeholder(doc, "Feedback", "Facilitator feedback form and Trainee feedback/grades panel")

    page_break(doc)
    add_page_title(doc, "13", "Notifications and Audit Logs")
    add_requirements(doc, [
        ("FR-069", "The system shall record supported business events in an audit log with actor, action, entity, message, module, metadata and timestamp.", "System", "Implemented - service"),
        ("FR-070", "The Admin dashboard shall list and filter audit activity by message, user, module and event attributes.", "Admin", "Implemented - UI/API data"),
        ("FR-071", "The system shall derive a user's notification feed from audit events and per-user read state.", "All", "Implemented - API"),
        ("FR-072", "The notification feed shall return an unread count with paginated notification items.", "All", "Implemented - API + UI"),
        ("FR-073", "A user shall be able to mark one notification as read.", "All", "Implemented - API + UI"),
        ("FR-074", "A user shall be able to mark all visible notifications as read.", "All", "Implemented - API + UI"),
        ("FR-075", "Notification read state shall be stored separately for each user and audit event.", "System", "Implemented - schema"),
    ], "audit.ts, notifications routes/service, AuditLog and NotificationRead schema, NotificationPanel/useNotifications and audit coverage tests.")
    add_gap_callout(doc, "Coverage boundary", "Audit logging covers users, batches, assignments, submissions, resources, sessions and feedback. Announcements and discussions do not have backend routes and therefore are not covered as persistent events.")
    add_screenshot_placeholder(doc, "Notifications and audit", "notification panel with unread states and Admin audit-log filters")

    page_break(doc)
    add_page_title(doc, "14", "Dashboards and Analytics")
    add_requirements(doc, [
        ("FR-076", "The Admin dashboard shall present system analytics, batch performance comparison, deadlines, recent activity and quick actions.", "Admin", "Implemented - UI"),
        ("FR-077", "The Facilitator dashboard shall present assigned batches, trainee counts, score/completion/attendance metrics, pending reviews and recent submissions.", "Facilitator", "Implemented - UI/API data"),
        ("FR-078", "The Facilitator batch detail page shall present trainee, assignment, attendance, session and submission summaries.", "Facilitator", "Implemented - UI/API data"),
        ("FR-079", "The Trainee dashboard shall present personal progress, assignments, sessions, resources, grades/feedback and facilitator contacts.", "Trainee", "Implemented - UI/API data"),
        ("FR-080", "Dashboard metrics shall display neutral/empty states when the required data is unavailable rather than inventing values.", "All", "Implemented - UI patterns"),
        ("FR-081", "The application shall provide stat cards, progress bars, status badges and simple bar charts for summarized operational data.", "All", "Implemented - components"),
    ], "role dashboard pages, FacilitatorBatchDetailPage.tsx, StatCard/ProgressBar/BarChart/EmptyState components and batch metrics service.")
    add_gap_callout(doc, "Meaning of real-time", "The Admin navigation uses the label 'Real-time Analytics', but the repository has no WebSocket or push layer. Values refresh through ordinary API/state updates.")
    add_screenshot_placeholder(doc, "Analytics", "Admin analytics dashboard and Facilitator/Trainee progress cards")

    page_break(doc)
    add_page_title(doc, "15", "Announcements and Reports")
    add_requirements(doc, [
        ("FR-082", "The frontend shall display announcements with title, message, audience, priority, pinned state, author/date and read count.", "All", "Frontend only"),
        ("FR-083", "Admin and Facilitator users shall be able to create an announcement within the current client session.", "Admin, Facilitator", "Frontend only"),
        ("FR-084", "The Trainee announcements view shall increment local read count when announcements are opened/viewed.", "Trainee", "Frontend only"),
        ("FR-085", "The Admin interface shall generate Attendance, Assignment, Performance, Feedback, Session, Resource Usage and Audit report datasets.", "Admin", "Implemented - client"),
        ("FR-086", "The Admin interface shall filter applicable report datasets by date range and export CSV files.", "Admin", "Implemented - client"),
        ("FR-087", "The Admin interface shall open a print-ready report view that can be saved as PDF through the browser print dialog.", "Admin", "Implemented - client"),
        ("FR-088", "The Admin interface shall allow Daily, Weekly or Monthly report schedules to be added and removed as local configuration entries.", "Admin", "Configuration only"),
    ], "announcements.service/store/mock data and role dashboards; AdminDashboardPage report builder/export/schedule functions.")
    add_gap_callout(doc, "Missing persistence", "No /api/announcements routes exist. Announcement posts and reads are not database-backed. Scheduled reports do not run because there is no backend scheduler/job runner.")
    add_screenshot_placeholder(doc, "Announcements and reports", "announcement card plus report export and scheduled-report configuration screens")

    page_break(doc)
    add_page_title(doc, "16", "Common User Experience Functions")
    add_requirements(doc, [
        ("FR-089", "The application shall route authenticated users to role-specific dashboards and redirect unauthorized role access.", "All", "Implemented - UI/API"),
        ("FR-090", "The dashboards shall provide consistent role-specific sidebar navigation and page headings.", "All", "Implemented - UI"),
        ("FR-091", "List screens shall provide relevant search, filter, sort, pagination and empty-state behavior.", "All", "Implemented - UI/API"),
        ("FR-092", "The UI shall use confirmation dialogs for destructive actions and saving indicators for asynchronous forms.", "All", "Implemented - UI"),
        ("FR-093", "The UI shall display success/error toast messages for completed or failed actions.", "All", "Implemented - UI"),
        ("FR-094", "The Admin interface shall provide global search across available client-side entities and navigation results.", "Admin", "Implemented - UI"),
        ("FR-095", "The application shall lazy-load major dashboard routes and show a loading fallback during route retrieval.", "All", "Implemented - App.tsx"),
        ("FR-096", "The application shall show a recoverable error fallback when an uncaught frontend render error occurs.", "All", "Implemented - ErrorBoundary"),
    ], "App.tsx, RequireAuth, navigation constants, reusable components, dashboard list screens, Toast and ErrorBoundary.")
    add_gap_callout(doc, "Contact/reminder behavior", "Several Contact or Reminder actions use mailto links or local audit/toast behavior; no internal messaging, SMS or email delivery workflow is implemented.")

    page_break(doc)
    add_page_title(doc, "17", "Business Rules - Access and Identity")
    add_rules(doc, [
        ("BR-001", "Only active, non-deleted accounts may authenticate or refresh a session.", "Auth service and middleware"),
        ("BR-002", "Login attempts are rate-limited; failed attempts contribute to temporary account lockout according to configuration.", "Rate limiter and auth service"),
        ("BR-003", "Passwords must be 8-72 characters and contain at least one letter and one number.", "Auth validator"),
        ("BR-004", "Only Admin may create invitations. An invite is valid only while Pending and before its expiry time.", "Auth route/service"),
        ("BR-005", "A raw invitation token is stored only as a hash and is not exposed in production unless explicitly configured.", "Auth service/config/tests"),
        ("BR-006", "Access-token protected routes require a Bearer token and re-check that the database user remains active.", "requireAuth middleware"),
        ("BR-007", "Frontend route checks do not replace backend authorization; both layers are present.", "App.tsx and route middleware"),
        ("BR-008", "Admin may act globally. Facilitator modification rights generally require object ownership or batch assignment. Trainee rights generally require ownership/enrollment.", "Service authorization rules"),
        ("BR-009", "Deleting users, batches, assignments, sessions or resources is implemented as soft deletion; records remain in storage for history.", "Prisma schema/services"),
        ("BR-010", "Deleting a user also deactivates the account and revokes active refresh tokens.", "Users service"),
    ])
    add_gap_callout(doc, "Exception", "The current forgot-password endpoint does not verify possession of an email token. BR-004 does not make that separate endpoint secure.")

    page_break(doc)
    add_page_title(doc, "18", "Business Rules - Training Operations")
    add_rules(doc, [
        ("BR-011", "Batch code must be unique. Supported programs are BA, Data Engineering, AI/ML and UI/UX; tracks are BTech and MBA.", "Schema and validators"),
        ("BR-012", "A trainee may be enrolled once per batch; removal is recorded with removedAt and a later enrollment can reactivate the row.", "BatchTrainee schema/service"),
        ("BR-013", "A multi-batch assignment must reference at least one batch and stores one primary batch for backward compatibility.", "Assignment validator/service"),
        ("BR-014", "There is at most one submission per assignment and trainee.", "Submission unique constraint"),
        ("BR-015", "Submission status is determined by timing and review state: Not Started, Under Review, Completed or Late.", "Submission enum/service"),
        ("BR-016", "Only the current attachment is presented by default; prior replacement files remain stored as superseded history.", "Submission attachment service"),
        ("BR-017", "Grades, when supplied, must be between 0 and 100.", "Submission validator"),
        ("BR-018", "Attendance is unique by session and trainee; status is Present, Absent, Late or Excused.", "Attendance schema/validator"),
        ("BR-019", "Feedback rating must be an integer from 1 to 5 and feedback cannot be edited or deleted through the API.", "Feedback validator/routes"),
        ("BR-020", "Resource update/delete requires uploader ownership unless the actor is Admin; each successful download increments the counter.", "Resources service"),
        ("BR-021", "Calendar access is based on Admin-global, Facilitator-managed or Trainee-enrolled batches.", "Calendar service/tests"),
        ("BR-022", "Audit-derived notification reads are per user and do not alter the underlying audit event.", "NotificationRead schema/service"),
    ])

    page_break(doc)
    add_page_title(doc, "19", "Non-Functional Requirements - Security")
    add_requirements(doc, [
        ("NFR-001", "All business API routes except documented public authentication/health endpoints shall require authentication.", "Security", "Implemented"),
        ("NFR-002", "Privileged API operations shall enforce role and ownership/enrollment checks on the server.", "Security", "Implemented; noted exceptions"),
        ("NFR-003", "Passwords shall be stored using bcrypt hashing and shall never be returned in API DTOs.", "Security", "Implemented"),
        ("NFR-004", "Refresh tokens and invite tokens shall be stored as hashes; refresh tokens shall support rotation and revocation.", "Security", "Implemented - tested"),
        ("NFR-005", "Production cookies shall be HttpOnly, Secure and suitable for cross-domain frontend/backend deployment.", "Security", "Implemented by config"),
        ("NFR-006", "The API shall use Helmet, explicit credentialed CORS configuration, request size limits and API/login/password-reset rate limits.", "Security", "Implemented"),
        ("NFR-007", "Uploaded files shall be validated by configured size/type controls and served through authorization-aware routes where private.", "Security", "Implemented; resource caveat"),
        ("NFR-008", "Production error responses shall not expose stack traces, connection strings or dependency internals.", "Privacy", "Implemented - tested"),
        ("NFR-009", "Secrets shall be supplied through environment variables and .env files shall not be committed.", "Security", "Configured/documented"),
        ("NFR-010", "The production password-recovery design should require a single-use expiring verification token before changing a password.", "Security", "Not implemented"),
    ], "app middleware, auth/storage configuration, validators, error handler, tests and deployment documentation.")

    page_break(doc)
    add_page_title(doc, "20", "Non-Functional Requirements - Performance and Reliability")
    add_requirements(doc, [
        ("NFR-011", "List endpoints shall use server-side pagination and database filtering/sorting for scalable result sets.", "Performance", "Implemented"),
        ("NFR-012", "Aggregate dashboard queries shall avoid per-row N+1 query patterns where tested.", "Performance", "Implemented - tested"),
        ("NFR-013", "The frontend shall lazy-load major role pages and produce optimized static assets through the Vite production build.", "Performance", "Implemented; build passed"),
        ("NFR-014", "The API shall use response compression and pooled PostgreSQL connections for hosted operation.", "Performance", "Implemented/configured"),
        ("NFR-015", "The service shall expose a liveness endpoint and a readiness endpoint that checks database and storage dependencies.", "Reliability", "Implemented - tested"),
        ("NFR-016", "Readiness failures shall return HTTP 503 with component-level ok/error status and no sensitive details.", "Reliability", "Implemented - tested"),
        ("NFR-017", "The backend shall perform graceful shutdown and disconnect Prisma on termination signals.", "Reliability", "Implemented"),
        ("NFR-018", "A failed database write after file upload shall remove the newly stored file to avoid orphans.", "Reliability", "Implemented - tested"),
        ("NFR-019", "Production uploads shall use persistent object storage; local filesystem storage is suitable only where the host volume is persistent.", "Reliability", "Configurable; deployment caveat"),
        ("NFR-020", "A live deployment shall be verified for database migration, restart persistence, storage persistence and cross-domain session behavior.", "Operations", "Not yet evidenced"),
    ], "pagination helpers/services, tests, Vite build, app/index, health route, storage providers and deployment guide.")

    page_break(doc)
    add_page_title(doc, "21", "Non-Functional Requirements - Usability and Maintainability")
    add_requirements(doc, [
        ("NFR-021", "The user interface shall use consistent role navigation, headings, colors, reusable controls and responsive layouts.", "Usability", "Implemented"),
        ("NFR-022", "Forms shall provide labels, validation feedback, saving state and cancellation/confirmation behavior.", "Usability", "Implemented"),
        ("NFR-023", "Data views shall provide loading skeletons, empty states, searchable/filterable lists and visible status indicators.", "Usability", "Implemented"),
        ("NFR-024", "Interactive dialogs and controls shall support keyboard escape/click-outside patterns and accessible labels where implemented.", "Accessibility", "Partially implemented"),
        ("NFR-025", "The application shall recover from uncaught frontend rendering failures with an error boundary and reload action.", "Usability", "Implemented"),
        ("NFR-026", "The codebase shall use shared TypeScript types, validators, services, components and stores to reduce duplication.", "Maintainability", "Implemented"),
        ("NFR-027", "The backend shall expose OpenAPI/Swagger documentation outside production when enabled.", "Maintainability", "Implemented"),
        ("NFR-028", "Database changes shall be versioned through Prisma migrations and validated before deployment.", "Maintainability", "Implemented; live apply pending"),
        ("NFR-029", "Automated tests shall cover authentication, authorization, validation, core CRUD, uploads, health, audit and query behavior.", "Quality", "92 tests pass"),
        ("NFR-030", "Before academic submission, the interface should be checked against WCAG keyboard, focus, contrast and screen-reader criteria.", "Accessibility", "Formal audit not found"),
    ], "frontend components/hooks/styles, TypeScript structure, Swagger config, Prisma migrations and backend test suite.")

    page_break(doc)
    add_page_title(doc, "22", "Data, Integrations and Constraints")
    doc.add_paragraph("Core records", style="Heading 2")
    add_definition_rows(doc, [
        ("Identity", "Role, Permission, User, UserProfile, UserInvite, RefreshToken"),
        ("Training", "Batch, BatchTrainee, Assignment, AssignmentBatch, Submission, SubmissionAttachment"),
        ("Delivery", "Session, Attendance, Resource, FeedbackEntry"),
        ("Engagement", "Announcement, AnnouncementRead, DiscussionThread, DiscussionMessage"),
        ("Operations", "AuditLog, NotificationRead"),
    ])
    doc.add_paragraph("Implemented integration boundaries", style="Heading 2")
    add_bullets(doc, [
        "PostgreSQL/Prisma for durable business data; migrations are present but were not proven against a reachable live database in this review.",
        "Local or S3 storage through a provider interface. Cloudinary is not implemented.",
        "Console email provider only. No real SES/Postmark/Resend connection exists.",
        "Browser mailto links are used for selected contact actions; no internal messaging service exists.",
        "No WebSocket, push notification, background job or scheduled-report service exists.",
        "Deployment files target Vercel plus Railway/Render and a Neon-compatible database, but the repository does not prove a completed live deployment.",
    ], compact=True)
    add_gap_callout(doc, "Dormant schema", "Announcement and discussion tables exist in Prisma, but corresponding backend routes/services are absent. Their presence in the schema must not be interpreted as completed functionality.")

    page_break(doc)
    add_page_title(doc, "23", "Known Gaps and Future Enhancements")
    doc.add_paragraph("Confirmed gaps in the current implementation", style="Heading 2")
    add_bullets(doc, [
        "Announcements and announcement reads are frontend-only and do not persist across a real restart/session.",
        "Scheduled reports are configuration-only; no backend scheduler creates or sends files.",
        "No real email provider sends invitations, password recovery, reminders or notifications.",
        "Forgot Password changes an active account password using only email and a new password.",
        "No WebSocket/push layer exists despite 'real-time' wording in the navigation.",
        "No backend discussion/messaging module is exposed; selected contact actions use mailto.",
        "Server-side least-privilege should be tightened for submission roster reads and trainee resource verification filtering.",
        "No evidence of a completed production deploy, live migration, backup restore or formal accessibility test was found.",
    ], compact=True)
    doc.add_paragraph("Prioritized future enhancements", style="Heading 2")
    add_definition_rows(doc, [
        ("P1 - Security", "Token-based password recovery; tighten submission/resource authorization; security review."),
        ("P1 - Persistence", "Implement announcement CRUD/read APIs and connect all role views to PostgreSQL."),
        ("P1 - Communication", "Integrate a real email provider for invitations and recovery."),
        ("P2 - Automation", "Add background jobs for scheduled reports and email delivery with run history."),
        ("P2 - Messaging", "Implement discussion/internal messaging APIs or remove dormant schema and UI expectations."),
        ("P2 - Operations", "Deploy to a live environment; test migrations, backups, S3 persistence and monitoring."),
        ("P3 - Experience", "Add push/live updates, formal WCAG testing, richer analytics and audit export controls."),
    ])
    add_screenshot_placeholder(doc, "Future evidence", "replace with final production deployment, monitoring and accessibility evidence when available")

    page_break(doc)
    add_page_title(doc, "24", "Conclusion and Verification Summary")
    add_text(doc, "The implemented Trainee Portal provides a substantial end-to-end foundation for academic demonstration: authenticated role-based access, batch and assignment operations, trainee submissions, grading, sessions, attendance, resources, feedback, audit activity, notifications and reporting views.")
    add_text(doc, "The strongest implementation areas are the typed API, relational data model, authorization tests, file-handling controls and role-specific user experience. The main delivery risks are not hidden: announcements and schedules are client-only, email and live updates are absent, password recovery is insecure for production, and a real hosted database/deployment has not been proven.")
    doc.add_paragraph("Verification performed for this FRD", style="Heading 2")
    add_definition_rows(doc, [
        ("Repository analysis", "Frontend pages/stores/services, backend routes/services/validators, Prisma schema/migrations, tests and deployment documentation reviewed."),
        ("Production build", "Passed for frontend (Vite/TypeScript) and backend (TypeScript) on 13 July 2026."),
        ("Automated tests", "20 backend test files passed; 92 tests passed."),
        ("Live environment", "Not verified; no claim is made that migrations, hosting, email or persistent object storage are operational."),
        ("Document scope", "96 functional requirements, 22 business rules and 30 non-functional requirements documented from the as-is implementation."),
    ])
    doc.add_paragraph("Recommended sign-off statement", style="Heading 2")
    p = doc.add_paragraph()
    border_paragraph(p, color="9FB3C8", fill=LIGHT_BLUE)
    r = p.add_run("This FRD accurately represents the reviewed implementation as of 13 July 2026, subject to the explicitly listed gaps and the absence of live deployment verification.")
    set_run_font(r, size=9.5, italic=True, color=NAVY)
    doc.add_paragraph("Final screenshot checklist", style="Heading 2")
    add_bullets(doc, [
        "Use consistent browser size and zoom; remove personal information.",
        "Capture login, each role dashboard, batch, assignment/submission, calendar, resource, feedback, audit and report screens.",
        "Replace placeholders only after final data and styling are stable.",
    ], compact=True)

    # Core metadata and deterministic document properties.
    props = doc.core_properties
    props.title = "Trainee Portal Functional Requirements Document"
    props.subject = "As-is functional requirements based on implemented source code"
    props.author = "Business Analyst"
    props.keywords = "FRD, Trainee Portal, functional requirements, business analysis"
    props.comments = "Generated from repository evidence; known gaps are explicitly identified."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
