from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "frd" / "Trainee_Portal_Functional_Requirements_Document_v1.0.docx"
PDF = ROOT / "output" / "pdf" / "Trainee_Portal_Functional_Requirements_Document_v1.0.pdf"

NAVY = colors.HexColor("#17365D")
BLUE = colors.HexColor("#2E74B5")
LIGHT_BLUE = colors.HexColor("#EAF2F8")
PALE = colors.HexColor("#F4F6F9")
MUTED = colors.HexColor("#5B6573")
GRID = colors.HexColor("#B8C4CE")


def register_fonts():
    candidates = [
        ("Calibri", Path("C:/Windows/Fonts/calibri.ttf")),
        ("Calibri-Bold", Path("C:/Windows/Fonts/calibrib.ttf")),
        ("Calibri-Italic", Path("C:/Windows/Fonts/calibrii.ttf")),
    ]
    for name, path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont(name, str(path)))
    return "Calibri" if "Calibri" in pdfmetrics.getRegisteredFontNames() else "Helvetica"


FONT = register_fonts()
BOLD = "Calibri-Bold" if "Calibri-Bold" in pdfmetrics.getRegisteredFontNames() else "Helvetica-Bold"
ITALIC = "Calibri-Italic" if "Calibri-Italic" in pdfmetrics.getRegisteredFontNames() else "Helvetica-Oblique"


def clean(text):
    return (
        text.replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2022", "-")
        .replace("\u00a0", " ")
    )


styles = getSampleStyleSheet()
BODY = ParagraphStyle("Body", fontName=FONT, fontSize=8.7, leading=10.2, textColor=colors.HexColor("#20252B"), spaceAfter=4)
H1 = ParagraphStyle("H1", fontName=BOLD, fontSize=16, leading=18, textColor=BLUE, spaceAfter=8, keepWithNext=True)
H2 = ParagraphStyle("H2", fontName=BOLD, fontSize=11, leading=13, textColor=NAVY, spaceBefore=5, spaceAfter=4, keepWithNext=True)
H3 = ParagraphStyle("H3", fontName=BOLD, fontSize=9.5, leading=11, textColor=NAVY, spaceBefore=4, spaceAfter=3, keepWithNext=True)
SMALL = ParagraphStyle("Small", fontName=FONT, fontSize=7.2, leading=8.3, textColor=MUTED)
CELL = ParagraphStyle("Cell", fontName=FONT, fontSize=6.9, leading=7.8, textColor=colors.HexColor("#20252B"))
CELL_BOLD = ParagraphStyle("CellBold", fontName=BOLD, fontSize=7.1, leading=8.0, textColor=NAVY)
CELL_HEAD = ParagraphStyle("CellHead", fontName=BOLD, fontSize=7.2, leading=8.1, textColor=colors.white, alignment=TA_LEFT)
TOC = ParagraphStyle("TOC", fontName=FONT, fontSize=8.4, leading=10.3, textColor=colors.HexColor("#20252B"), leftIndent=4)
PLACEHOLDER = ParagraphStyle("Placeholder", fontName=FONT, fontSize=7.8, leading=9.5, textColor=MUTED, borderColor=GRID, borderWidth=0.6, borderPadding=8, backColor=colors.HexColor("#F7F9FB"), spaceBefore=11, spaceAfter=8)
CALLOUT = ParagraphStyle("Callout", fontName=FONT, fontSize=7.8, leading=9.4, textColor=colors.HexColor("#5A4300"), borderColor=colors.HexColor("#D5A021"), borderWidth=0.6, borderPadding=7, backColor=colors.HexColor("#FFF8E6"), spaceBefore=11, spaceAfter=9)


def has_page_break(paragraph):
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def paragraph_flowable(p):
    text = clean(p.text.strip())
    if not text:
        return Spacer(1, 2)
    style_name = p.style.name if p.style else "Normal"
    if style_name == "Heading 1":
        return Paragraph(escape(text), H1)
    if style_name == "Heading 2":
        return Paragraph(escape(text), H2)
    if style_name == "Heading 3":
        return Paragraph(escape(text), H3)
    if style_name.startswith("List Bullet"):
        return Paragraph(f"- {escape(text)}", ParagraphStyle("Bullet", parent=BODY, leftIndent=13, firstLineIndent=-8, spaceAfter=2))
    if text.startswith("SCREENSHOT PLACEHOLDER"):
        lines = text.splitlines()
        html = f"<b><font color='#2E74B5'>{escape(lines[0])}</font></b>"
        if len(lines) > 1:
            html += "<br/>" + "<br/>".join(escape(x) for x in lines[1:])
        return Paragraph(html, PLACEHOLDER)
    if text.startswith(("Known security gap:", "Scope caution:", "Interpretation:", "Access caveat:", "UI placeholder:", "Demo limitation:", "Visibility rule caveat:", "Coverage boundary:", "Meaning of real-time:", "Missing persistence:", "Contact/reminder behavior:", "Exception:", "Dormant schema:")):
        head, _, rest = text.partition(":")
        return Paragraph(f"<b>{escape(head)}:</b>{escape(rest)}", CALLOUT)
    if text.startswith("Document navigation note:"):
        return Paragraph(f"<b>{escape(text)}</b>", ParagraphStyle("NavNote", parent=SMALL, backColor=LIGHT_BLUE, borderColor=GRID, borderWidth=0.5, borderPadding=6, spaceBefore=10))
    # Retain the visible bold lead used by purpose-like paragraphs.
    for lead in ("Purpose:",):
        if text.startswith(lead):
            return Paragraph(f"<b><font color='#17365D'>{escape(lead)}</font></b>{escape(text[len(lead):])}", BODY)
    return Paragraph(escape(text), BODY)


def table_flowable(tbl):
    rows = []
    for r_idx, row in enumerate(tbl.rows):
        rendered = []
        for c_idx, cell in enumerate(row.cells):
            txt = clean("\n".join(p.text for p in cell.paragraphs).strip())
            if r_idx == 0:
                st = CELL_HEAD
            elif c_idx == 0:
                st = CELL_BOLD
            else:
                st = CELL
            rendered.append(Paragraph(escape(txt).replace("\n", "<br/>"), st))
        rows.append(rendered)
    cols = len(rows[0]) if rows else 1
    if cols == 4:
        widths = [0.52 * inch, 3.70 * inch, 0.80 * inch, 1.48 * inch]
    elif cols == 3:
        widths = [0.62 * inch, 4.02 * inch, 1.86 * inch]
    elif cols == 2:
        widths = [1.52 * inch, 4.98 * inch]
    else:
        widths = [6.5 * inch / cols] * cols
    t = Table(rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
    commands = [
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("GRID", (0, 0), (-1, -1), 0.45, GRID),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    if len(rows) > 1:
        commands.append(("BACKGROUND", (0, 1), (0, -1), LIGHT_BLUE))
    t.setStyle(TableStyle(commands))
    return t


def header_footer(canvas, doc):
    canvas.saveState()
    page = canvas.getPageNumber()
    if page > 1:
        canvas.setFont(BOLD, 7.2)
        canvas.setFillColor(MUTED)
        canvas.drawString(0.85 * inch, 10.64 * inch, "TRAINEE PORTAL  |  FUNCTIONAL REQUIREMENTS DOCUMENT")
    canvas.setFont(FONT, 7.5)
    canvas.setFillColor(MUTED)
    canvas.drawRightString(7.65 * inch, 0.34 * inch, f"Page {page}")
    canvas.restoreState()


def build_pdf():
    source = Document(DOCX)
    doc = BaseDocTemplate(
        str(PDF),
        pagesize=letter,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.67 * inch,
        bottomMargin=0.58 * inch,
        title="Trainee Portal Functional Requirements Document",
        author="Business Analyst",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates(PageTemplate(id="frd", frames=[frame], onPage=header_footer))

    story = []
    body = source.element.body
    para_by_el = {p._p: p for p in source.paragraphs}
    table_by_el = {t._tbl: t for t in source.tables}
    first_content = True
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_by_el.get(child)
            if p is None:
                continue
            if has_page_break(p):
                story.append(PageBreak())
                first_content = False
                continue
            text = p.text.strip()
            if first_content and text == "":
                story.append(Spacer(1, 12))
                continue
            # Special treatment for the cover title and subtitle.
            if first_content and text == "FUNCTIONAL REQUIREMENTS\nDOCUMENT":
                story.append(Spacer(1, 1.05 * inch))
                story.append(Paragraph("FUNCTIONAL REQUIREMENTS<br/>DOCUMENT", ParagraphStyle("CoverTitle", fontName=BOLD, fontSize=28, leading=31, alignment=TA_CENTER, textColor=NAVY, spaceAfter=12)))
                continue
            if first_content and text == "Trainee Portal":
                story.append(Paragraph(text, ParagraphStyle("CoverProject", fontName=BOLD, fontSize=20, leading=23, alignment=TA_CENTER, textColor=BLUE, spaceAfter=18)))
                continue
            if first_content and text == "As-Is Functional Specification Based on Implemented Source Code":
                story.append(Paragraph(text, ParagraphStyle("CoverSub", fontName=ITALIC, fontSize=10.5, leading=13, alignment=TA_CENTER, textColor=MUTED, spaceAfter=32)))
                continue
            if first_content and text.startswith("Prepared from the repository"):
                story.append(Spacer(1, 24))
                story.append(Paragraph(escape(text), ParagraphStyle("CoverNote", fontName=FONT, fontSize=8, leading=10, alignment=TA_CENTER, textColor=MUTED)))
                continue
            if text == "Table of Contents":
                story.append(Paragraph(text, H1))
                continue
            if p.text.count("\t") == 1 and p.text.strip().split("\t")[-1].isdigit():
                left, page = clean(p.text).split("\t")
                data = [[Paragraph(escape(left), TOC), Paragraph(escape(page), ParagraphStyle("TOCPage", parent=TOC, alignment=TA_RIGHT))]]
                toc_row = Table(data, colWidths=[6.1 * inch, 0.4 * inch], hAlign="LEFT")
                toc_row.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("BOTTOMPADDING", (0, 0), (-1, -1), 1), ("TOPPADDING", (0, 0), (-1, -1), 1)]))
                story.append(toc_row)
                continue
            story.append(paragraph_flowable(p))
        elif child.tag == qn("w:tbl"):
            tbl = table_by_el.get(child)
            if tbl is not None:
                story.append(table_flowable(tbl))
                story.append(Spacer(1, 4))
    doc.build(story)
    pages = len(PdfReader(str(PDF)).pages)
    print(f"{PDF}\nPages={pages}")


if __name__ == "__main__":
    build_pdf()
