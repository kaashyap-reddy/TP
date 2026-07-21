from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "frd" / "Trainee_Portal_Functional_Requirements_Document_v2.0.docx"
PDF = ROOT / "output" / "pdf" / "Trainee_Portal_Functional_Requirements_Document_v2.0.pdf"

NAVY = colors.HexColor("#17365D")
BLUE = colors.HexColor("#2E74B5")
LIGHT_BLUE = colors.HexColor("#EAF2F8")
PALE = colors.HexColor("#F4F6F9")
MUTED = colors.HexColor("#5B6573")
GRID = colors.HexColor("#B8C4CE")
GREEN_FILL = colors.HexColor("#E8F5EE")
GREEN_BORDER = colors.HexColor("#7FB39B")
GREEN_TEXT = colors.HexColor("#1F4A38")
AMBER_FILL = colors.HexColor("#FFF8E6")
AMBER_BORDER = colors.HexColor("#D5A021")
AMBER_TEXT = colors.HexColor("#5A4300")


def register_fonts():
    candidates = [
        ("Calibri", Path("C:/Windows/Fonts/calibri.ttf")),
        ("Calibri-Bold", Path("C:/Windows/Fonts/calibrib.ttf")),
        ("Calibri-Italic", Path("C:/Windows/Fonts/calibrii.ttf")),
    ]
    for name, path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont(name, str(path)))
    return "Calibri" if "Calibri" in pdfmetrics.getRegisteredFontNames() else "Helvetica"


FONT = register_fonts()
BOLD = "Calibri-Bold" if "Calibri-Bold" in pdfmetrics.getRegisteredFontNames() else "Helvetica-Bold"
ITALIC = "Calibri-Italic" if "Calibri-Italic" in pdfmetrics.getRegisteredFontNames() else "Helvetica-Oblique"


def clean(text):
    return (
        text.replace("–", "-")
        .replace("—", "-")
        .replace("‘", "'")
        .replace("’", "'")
        .replace("“", '"')
        .replace("”", '"')
        .replace("•", "-")
        .replace(" ", " ")
    )


styles = getSampleStyleSheet()
BODY = ParagraphStyle("Body", fontName=FONT, fontSize=8.7, leading=10.2, textColor=colors.HexColor("#20252B"), spaceAfter=4)
H1 = ParagraphStyle("H1", fontName=BOLD, fontSize=16, leading=18, textColor=BLUE, spaceAfter=8, keepWithNext=True)
H2 = ParagraphStyle("H2", fontName=BOLD, fontSize=11, leading=13, textColor=NAVY, spaceBefore=5, spaceAfter=4, keepWithNext=True)
H3 = ParagraphStyle("H3", fontName=BOLD, fontSize=9.5, leading=11, textColor=NAVY, spaceBefore=4, spaceAfter=3, keepWithNext=True)
SUB = ParagraphStyle("Sub", fontName=ITALIC, fontSize=8.3, leading=10, textColor=MUTED, spaceAfter=6)
SMALL = ParagraphStyle("Small", fontName=FONT, fontSize=7.2, leading=8.3, textColor=MUTED)
CELL = ParagraphStyle("Cell", fontName=FONT, fontSize=6.9, leading=7.8, textColor=colors.HexColor("#20252B"))
CELL_BOLD = ParagraphStyle("CellBold", fontName=BOLD, fontSize=7.1, leading=8.0, textColor=NAVY)
CELL_HEAD = ParagraphStyle("CellHead", fontName=BOLD, fontSize=7.2, leading=8.1, textColor=colors.white, alignment=TA_LEFT)
TOC = ParagraphStyle("TOC", fontName=FONT, fontSize=8.4, leading=10.3, textColor=colors.HexColor("#20252B"), leftIndent=4)
TOC_BOLD = ParagraphStyle("TOCBold", parent=TOC, fontName=BOLD, textColor=NAVY)
PLACEHOLDER = ParagraphStyle("Placeholder", fontName=FONT, fontSize=7.8, leading=9.5, textColor=MUTED, borderColor=GRID, borderWidth=0.6, borderPadding=8, backColor=colors.HexColor("#F7F9FB"), spaceBefore=11, spaceAfter=8)
CALLOUT_AMBER = ParagraphStyle("CalloutAmber", fontName=FONT, fontSize=7.8, leading=9.4, textColor=AMBER_TEXT, borderColor=AMBER_BORDER, borderWidth=0.6, borderPadding=7, backColor=AMBER_FILL, spaceBefore=11, spaceAfter=9)
CALLOUT_GREEN = ParagraphStyle("CalloutGreen", fontName=FONT, fontSize=7.8, leading=9.4, textColor=GREEN_TEXT, borderColor=GREEN_BORDER, borderWidth=0.6, borderPadding=7, backColor=GREEN_FILL, spaceBefore=11, spaceAfter=9)
CALLOUT_BLUE = ParagraphStyle("CalloutBlue", fontName=FONT, fontSize=7.8, leading=9.4, textColor=NAVY, borderColor=colors.HexColor("#C6D4E1"), borderWidth=0.6, borderPadding=7, backColor=LIGHT_BLUE, spaceBefore=10, spaceAfter=8)


def get_shading_fill(paragraph):
    """Read the actual w:shd fill color python-docx's border_paragraph() wrote onto this
    paragraph, so callout/placeholder/note boxes render correctly in the PDF regardless of
    what text they contain — no hardcoded prefix list needed."""
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return (shd.get(qn("w:fill")) or "").upper()


def has_page_break(paragraph):
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def paragraph_flowable(p):
    text = clean(p.text.strip())
    if not text:
        return Spacer(1, 2)
    style_name = p.style.name if p.style else "Normal"
    if style_name == "Heading 1":
        return Paragraph(escape(text), H1)
    if style_name == "Heading 2":
        return Paragraph(escape(text), H2)
    if style_name == "Heading 3":
        return Paragraph(escape(text), H3)
    if style_name.startswith("List Bullet"):
        return Paragraph(f"- {escape(text)}", ParagraphStyle("Bullet", parent=BODY, leftIndent=13, firstLineIndent=-8, spaceAfter=2))
    if text.startswith("SCREENSHOT PLACEHOLDER"):
        lines = text.splitlines()
        html = f"<b><font color='#2E74B5'>{escape(lines[0])}</font></b>"
        if len(lines) > 1:
            html += "<br/>" + "<br/>".join(escape(x) for x in lines[1:])
        return Paragraph(html, PLACEHOLDER)

    fill = get_shading_fill(p)
    if fill == "FFF8E6":  # amber gap/caution callout
        head, sep, rest = text.partition(": ")
        html = f"<b>{escape(head)}{escape(sep)}</b>{escape(rest)}" if sep else escape(text)
        return Paragraph(html, CALLOUT_AMBER)
    if fill == "E8F5EE":  # green "resolved since v1.0" callout
        head, sep, rest = text.partition(": ")
        html = f"<b>{escape(head)}{escape(sep)}</b>{escape(rest)}" if sep else escape(text)
        return Paragraph(html, CALLOUT_GREEN)
    if fill == "EAF2F8":  # light-blue note / sign-off box
        head, sep, rest = text.partition(": ")
        html = f"<b>{escape(head)}{escape(sep)}</b>{escape(rest)}" if sep else f"<i>{escape(text)}</i>"
        return Paragraph(html, CALLOUT_BLUE)

    # Section subtitle line immediately following a Heading 1 (italic, muted) — heuristically
    # any short italic-run-only paragraph is treated as body; subtitles are written as plain
    # paragraphs with italic runs, which python-docx round-trips as plain text here, so just
    # fall through to BODY — acceptable fidelity loss (italic emphasis, not content).
    for lead in ("Purpose:",):
        if text.startswith(lead):
            return Paragraph(f"<b><font color='#17365D'>{escape(lead)}</font></b>{escape(text[len(lead):])}", BODY)
    return Paragraph(escape(text), BODY)


def table_flowable(tbl):
    rows = []
    for r_idx, row in enumerate(tbl.rows):
        rendered = []
        for c_idx, cell in enumerate(row.cells):
            txt = clean("\n".join(p.text for p in cell.paragraphs).strip())
            if r_idx == 0:
                st = CELL_HEAD
            elif c_idx == 0:
                st = CELL_BOLD
            else:
                st = CELL
            rendered.append(Paragraph(escape(txt).replace("\n", "<br/>"), st))
        rows.append(rendered)
    cols = len(rows[0]) if rows else 1

    # Read the actual column widths set in the docx (set_table_geometry's gridCol values, in
    # twips) rather than guessing by column count — different 3-column tables (e.g. Rule/
    # Statement/Source vs. Endpoint/Access/Summary) need very different proportions, and a
    # count-only guess squeezes long text (e.g. "POST /auth/invite/accept") into a too-narrow
    # column and forces an ugly mid-word wrap.
    grid_widths = [int(col.get(qn("w:w"))) for col in tbl._tbl.tblGrid.findall(qn("w:gridCol"))]
    if grid_widths and len(grid_widths) == cols and sum(grid_widths) > 0:
        total_target = 6.5 * inch
        scale = total_target / sum(grid_widths)
        widths = [w * scale for w in grid_widths]
    elif cols == 4:
        widths = [0.52 * inch, 3.70 * inch, 0.80 * inch, 1.48 * inch]
    elif cols == 3:
        widths = [0.62 * inch, 4.02 * inch, 1.86 * inch]
    elif cols == 2:
        widths = [1.52 * inch, 4.98 * inch]
    else:
        widths = [6.5 * inch / cols] * cols
    t = Table(rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
    commands = [
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("GRID", (0, 0), (-1, -1), 0.45, GRID),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    if len(rows) > 1:
        commands.append(("BACKGROUND", (0, 1), (0, -1), LIGHT_BLUE))
    t.setStyle(TableStyle(commands))
    return t


def header_footer(canvas, doc):
    canvas.saveState()
    page = canvas.getPageNumber()
    if page > 1:
        canvas.setFont(BOLD, 7.2)
        canvas.setFillColor(MUTED)
        canvas.drawString(0.85 * inch, 10.64 * inch, "TRAINEE PORTAL  |  FUNCTIONAL REQUIREMENTS DOCUMENT  |  v2.0")
    canvas.setFont(FONT, 7.5)
    canvas.setFillColor(MUTED)
    canvas.drawRightString(7.65 * inch, 0.34 * inch, f"Page {page}")
    canvas.restoreState()


def build_pdf():
    source = Document(DOCX)
    PDF.parent.mkdir(parents=True, exist_ok=True)
    doc = BaseDocTemplate(
        str(PDF),
        pagesize=letter,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.67 * inch,
        bottomMargin=0.58 * inch,
        title="Trainee Portal Functional Requirements Document v2.0",
        author="Business Analyst",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates(PageTemplate(id="frd", frames=[frame], onPage=header_footer))

    story = []
    body = source.element.body
    para_by_el = {p._p: p for p in source.paragraphs}
    table_by_el = {t._tbl: t for t in source.tables}
    first_content = True
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_by_el.get(child)
            if p is None:
                continue
            if has_page_break(p):
                story.append(PageBreak())
                first_content = False
                continue
            text = p.text.strip()
            if first_content and text == "":
                story.append(Spacer(1, 12))
                continue
            if first_content and text == "FUNCTIONAL REQUIREMENTS\nDOCUMENT":
                story.append(Spacer(1, 1.05 * inch))
                story.append(Paragraph("FUNCTIONAL REQUIREMENTS<br/>DOCUMENT", ParagraphStyle("CoverTitle", fontName=BOLD, fontSize=28, leading=31, alignment=TA_CENTER, textColor=NAVY, spaceAfter=12)))
                continue
            if first_content and text == "Trainee Portal":
                story.append(Paragraph(text, ParagraphStyle("CoverProject", fontName=BOLD, fontSize=20, leading=23, alignment=TA_CENTER, textColor=BLUE, spaceAfter=18)))
                continue
            if first_content and text == "As-Is Functional Specification Based on Implemented Source Code":
                story.append(Paragraph(text, ParagraphStyle("CoverSub", fontName=ITALIC, fontSize=10.5, leading=13, alignment=TA_CENTER, textColor=MUTED, spaceAfter=32)))
                continue
            if first_content and text.startswith("Prepared from the repository"):
                story.append(Spacer(1, 24))
                story.append(Paragraph(escape(text), ParagraphStyle("CoverNote", fontName=FONT, fontSize=8, leading=10, alignment=TA_CENTER, textColor=MUTED)))
                continue
            if text == "Table of Contents":
                story.append(Paragraph(text, H1))
                continue
            if p.text.count("\t") == 1 and p.text.strip().split("\t")[-1].isdigit():
                left, pagenum = clean(p.text).split("\t")
                is_bold = any(run.bold for run in p.runs)
                style = TOC_BOLD if is_bold else TOC
                page_style = ParagraphStyle("TOCPage", parent=style, alignment=TA_RIGHT)
                data = [[Paragraph(escape(left), style), Paragraph(escape(pagenum), page_style)]]
                toc_row = Table(data, colWidths=[6.1 * inch, 0.4 * inch], hAlign="LEFT")
                toc_row.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("BOTTOMPADDING", (0, 0), (-1, -1), 1), ("TOPPADDING", (0, 0), (-1, -1), 1)]))
                story.append(toc_row)
                continue
            story.append(paragraph_flowable(p))
        elif child.tag == qn("w:tbl"):
            tbl = table_by_el.get(child)
            if tbl is not None:
                story.append(table_flowable(tbl))
                story.append(Spacer(1, 4))
    doc.build(story)
    pages = len(PdfReader(str(PDF)).pages)
    print(f"{PDF}\nPages={pages}")


if __name__ == "__main__":
    build_pdf()
