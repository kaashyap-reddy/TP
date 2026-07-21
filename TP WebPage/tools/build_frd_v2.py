from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "frd"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "Trainee_Portal_Functional_Requirements_Document_v2.0.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
PALE = "F4F6F9"
GRID = "B8C4CE"
MUTED = "5B6573"
GREEN = "2E7D32"
AMBER = "9A6700"
RED = "9B1C1C"
WHITE = "FFFFFF"


# ---------------------------------------------------------------------------
# Low-level formatting helpers (same conventions as v1.0's build_frd.py)
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def border_paragraph(paragraph, color=GRID, fill=PALE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), color)
        p_bdr.append(el)
    p_pr.append(p_bdr)


def add_field_toc_note(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    border_paragraph(p, color="C6D4E1", fill=LIGHT_BLUE)
    r = p.add_run("Document navigation note: ")
    set_run_font(r, size=8.5, bold=True, color=NAVY)
    r = p.add_run("The table of contents is static for reliable PDF export. Page numbers match the verified final rendering.")
    set_run_font(r, size=8.5, color=MUTED)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string("20252B")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 10),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 17, BLUE, 0, 9),
        ("Heading 2", 12.5, NAVY, 9, 5),
        ("Heading 3", 10.5, NAVY, 6, 3),
    ):
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = style_name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("TRAINEE PORTAL  |  FUNCTIONAL REQUIREMENTS DOCUMENT  |  v2.0")
    set_run_font(r, size=8, bold=True, color=MUTED)
    fp = section.footer.paragraphs[0]
    add_page_number(fp)
    return doc


def page_break(doc):
    doc.add_page_break()


def add_page_title(doc, number, title, subtitle=None):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.add_run(f"{number}. {title}" if number else title)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(subtitle)
        set_run_font(r2, size=9, italic=True, color=MUTED)


def add_text(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items, compact=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2 if compact else 4)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(item)
        set_run_font(r, size=9.5 if compact else 10)


def add_definition_rows(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Definition / Value"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9, bold=True, color=WHITE)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    set_run_font(run, size=9, bold=(idx == 0), color=NAVY if idx == 0 else None)
    set_table_geometry(table, [2200, 7160])
    return table


def add_requirements(doc, requirements, source_note=None):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["ID", "Functional requirement", "User(s)", "Implementation evidence"]
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = label
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8.3, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for req_id, req, users, evidence in requirements:
        cells = table.add_row().cells
        values = [req_id, req, users, evidence]
        for idx, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            if idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=7.8 if idx in (0, 2, 3) else 8.2, bold=(idx == 0), color=NAVY if idx == 0 else None)
    set_table_geometry(table, [780, 5330, 1150, 2100])
    if source_note:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Source basis: ")
        set_run_font(r, size=7.6, bold=True, color=MUTED)
        r = p.add_run(source_note)
        set_run_font(r, size=7.6, color=MUTED)
    return table


def add_rules(doc, rules, title="Business rules"):
    if title:
        doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, label in enumerate(("Rule", "Statement", "Source / rationale")):
        cell = table.rows[0].cells[i]
        cell.text = label
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for rid, statement, source in rules:
        cells = table.add_row().cells
        for idx, value in enumerate((rid, statement, source)):
            cells[idx].text = value
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT_BLUE)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=8.1, bold=(idx == 0), color=NAVY if idx == 0 else None)
    set_table_geometry(table, [900, 5780, 2680])


def add_screenshot_placeholder(doc, label, suggested_capture):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    border_paragraph(p, color="AAB7C4", fill="F7F9FB")
    r = p.add_run(f"SCREENSHOT PLACEHOLDER - {label}\n")
    set_run_font(r, size=9, bold=True, color=BLUE)
    r = p.add_run(f"Suggested capture: {suggested_capture}\nInsert final application screenshot here; retain sensitive/demo data only.")
    set_run_font(r, size=8.3, color=MUTED)


def add_gap_callout(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    border_paragraph(p, color="D5A021", fill="FFF8E6")
    r = p.add_run(f"{title}: ")
    set_run_font(r, size=9, bold=True, color=AMBER)
    r = p.add_run(text)
    set_run_font(r, size=9, color="5A4300")


def add_resolved_callout(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    border_paragraph(p, color="7FB39B", fill="E8F5EE")
    r = p.add_run(f"{title}: ")
    set_run_font(r, size=9, bold=True, color=GREEN)
    r = p.add_run(text)
    set_run_font(r, size=9, color="1F4A38")


def add_model_table(doc, model_name, purpose, fields):
    doc.add_paragraph(model_name, style="Heading 3")
    if purpose:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(purpose)
        set_run_font(r, size=8.6, italic=True, color=MUTED)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field(s)"
    hdr[1].text = "Purpose / notes"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8, bold=True, color=WHITE)
    for field, note in fields:
        cells = table.add_row().cells
        cells[0].text = field
        cells[1].text = note
        set_cell_shading(cells[0], LIGHT_BLUE)
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=7.9, bold=(idx == 0), color=NAVY if idx == 0 else None)
    set_table_geometry(table, [2400, 6960])


def add_endpoint_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Endpoint", "Access", "Summary"]
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = label
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for endpoint, access, summary in rows:
        cells = table.add_row().cells
        for idx, value in enumerate((endpoint, access, summary)):
            cells[idx].text = value
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT_BLUE)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=7.6, bold=(idx == 0), color=NAVY if idx == 0 else None)
    set_table_geometry(table, [2500, 1550, 5310])


# ---------------------------------------------------------------------------
# Cover / control / TOC
# ---------------------------------------------------------------------------

def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FUNCTIONAL REQUIREMENTS\nDOCUMENT")
    set_run_font(r, size=29, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Trainee Portal")
    set_run_font(r, size=20, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("As-Is Functional Specification Based on Implemented Source Code")
    set_run_font(r, size=11, italic=True, color=MUTED)
    p.paragraph_format.space_after = Pt(62)
    add_definition_rows(doc, [
        ("Document version", "2.0 (supersedes v1.0, dated 13 July 2026)"),
        ("Prepared as", "Business Analyst deliverable - internship/final-year project"),
        ("Prepared date", "20 July 2026"),
        ("Project", "Trainee Portal Web Application"),
        ("Confidentiality", "Academic / internal project use"),
    ])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    r = p.add_run("Prepared from the repository implementation; no unverified features have been added.")
    set_run_font(r, size=9, color=MUTED)


def add_control_page(doc):
    p = doc.add_paragraph("Document Control", style="Heading 1")
    doc.add_paragraph("Version history", style="Heading 2")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(("Version", "Date", "Author / Role", "Change summary")):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], NAVY)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_run_font(run, size=9, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("0.1", "13 Jul 2026", "Business Analyst", "Initial implementation review and requirement inventory."),
        ("1.0", "13 Jul 2026", "Business Analyst", "First FRD: 96 functional requirements, 22 business rules, 30 NFRs, 24 sections / 27 pages."),
        ("2.0", "20 Jul 2026", "Business Analyst", "Full re-audit against the current repository. Adds the Training Plan / curriculum-template module, "
         "Session Feedback and Assignment Feedback form-link features, real database-backed Announcements, "
         "token-verified Forgot Password, and the Sentry observability hook points - none of which existed at v1.0. "
         "Expands requirement/rule/NFR coverage, adds a full data dictionary and API endpoint appendix, and "
         "re-verifies every build/test claim first-party against the current codebase."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs:
                set_run_font(run, size=8.5)
    set_table_geometry(table, [850, 1250, 1750, 5860])

    doc.add_paragraph("What changed between v1.0 and v2.0", style="Heading 2")
    add_bullets(doc, [
        "New module documented: Training Plans - reusable curriculum templates (sessions, assignments, resources, default announcements) that are copied into an independent, editable set of records for every new batch (Section 7).",
        "New module documented: Session Feedback and Assignment Feedback external form-link features, previously undocumented as a distinct capability (Sections 14-15).",
        "Announcements reclassified from Frontend only to fully implemented: /api/announcements now persists to PostgreSQL with per-user read tracking and full Demo Mode parity (Section 17).",
        "Forgot Password reclassified from a known security gap to implemented: it now issues a single-use, one-hour, hashed reset token instead of accepting a bare email + new password (Section 5).",
        "New NFR section: Observability and Monitoring, covering the Sentry hook points added to both frontend and backend since v1.0 (Section 29).",
        "Batch creation rewritten to reflect the current workflow: an admin now picks a Training Plan (not a raw program/track), and program/track/end date are all server-derived, not client-supplied (Section 8).",
        "New appendices: a full Prisma data dictionary (Appendix A) and a complete REST endpoint reference covering all 89 API routes (Appendix B), neither of which existed in v1.0.",
        "All test, build and pass/fail figures in this version were re-executed against the current repository on the preparation date, not carried forward from v1.0 or from interim notes.",
    ], compact=True)

    doc.add_paragraph("Review and approval", style="Heading 2")
    add_definition_rows(doc, [
        ("Prepared by", "____________________________  Date: __________"),
        ("Reviewed by", "____________________________  Date: __________"),
        ("Approved by", "____________________________  Date: __________"),
    ])
    doc.add_paragraph("Document conventions", style="Heading 2")
    add_bullets(doc, [
        "Implemented - supported by frontend/backend code and, where applicable, automated tests.",
        "Frontend only - visible behavior exists, but there is no persistent backend endpoint.",
        "Configuration only - the screen stores a setting locally but no scheduler/job executes it.",
        "Known gap - explicitly absent or incomplete in the current repository.",
        "A green callout marks something resolved since v1.0; an amber callout marks an open gap or caution.",
    ], compact=True)
    add_gap_callout(doc, "Scope caution", "This FRD specifies the current codebase. It certifies a verified local development environment (local PostgreSQL, connected and exercised end-to-end) but does not certify a live production deployment, a hosted database, or a connected email provider.")


TOC_ENTRIES = [
    ("1", "Introduction", 4),
    ("2", "Project Overview and Architecture", 5),
    ("3", "User Roles and Access", 6),
    ("4", "Information Architecture and Navigation", 7),
    ("5", "Authentication and Account Access", 8),
    ("6", "User and Profile Management", 9),
    ("7", "Training Plans and Curriculum Templates", 10),
    ("8", "Batch Management", 11),
    ("9", "Assignment Management", 12),
    ("10", "Submission and Grading", 13),
    ("11", "Sessions and Calendar", 14),
    ("12", "Attendance Management", 15),
    ("13", "Resource Library", 16),
    ("14", "Session Feedback (Form Links)", 17),
    ("15", "Assignment Feedback (Form Links)", 18),
    ("16", "Facilitator-Trainee Performance Feedback", 19),
    ("17", "Announcements", 20),
    ("18", "Notifications and Audit Logs", 21),
    ("19", "Dashboards and Analytics", 22),
    ("20", "Reports", 23),
    ("21", "Common User Experience Functions", 24),
    ("22", "Discussions (Dormant Schema)", 25),
    ("23", "Business Rules - Access and Identity", 26),
    ("24", "Business Rules - Training Plan and Curriculum Generation", 27),
    ("25", "Business Rules - Training Operations", 28),
    ("26", "Business Rules - Feedback and Communication", 29),
    ("27", "Non-Functional Requirements - Security", 30),
    ("28", "Non-Functional Requirements - Performance and Reliability", 31),
    ("29", "Non-Functional Requirements - Observability and Monitoring", 32),
    ("30", "Non-Functional Requirements - Usability and Accessibility", 33),
    ("31", "Non-Functional Requirements - Maintainability and Quality", 34),
    ("32", "Data, Integrations and Constraints", 35),
    ("33", "Known Gaps and Future Enhancements", 36),
    ("34", "Conclusion and Verification Summary", 37),
    ("A", "Appendix A - Data Dictionary", 38),
    ("B", "Appendix B - API Endpoint Reference", 42),
    ("C", "Appendix C - Environment Configuration Reference", 46),
    ("D", "Appendix D - Glossary", 47),
]


def add_toc(doc):
    doc.add_paragraph("Table of Contents", style="Heading 1")
    for num, title, page in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        bold = num in {"1", "5", "23", "27", "33", "34", "A", "B"}
        r = p.add_run(f"{num}. {title}\t{page}")
        set_run_font(r, size=9.2, bold=bold, color=NAVY if bold else None)
    add_field_toc_note(doc)


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    doc = setup_document()
    add_cover(doc)
    page_break(doc)
    add_control_page(doc)
    page_break(doc)
    add_toc(doc)

    # ---------------- 1. Introduction ----------------
    page_break(doc)
    add_page_title(doc, "1", "Introduction")
    add_text(doc, "Purpose: This Functional Requirements Document (FRD) defines what the Trainee Portal currently does, who can use each capability, and the principal rules that govern the implemented workflows. It supersedes v1.0 (13 July 2026), which was written before the Training Plan module, the Session/Assignment Feedback features, database-backed Announcements and token-verified password reset existed in the repository.", "Purpose:")
    add_text(doc, "This is an as-is specification derived from the repository, not a wish list. Every functional requirement, business rule and non-functional requirement below cites the code path that implements it; every gap is stated as a gap rather than implied as complete.")
    doc.add_paragraph("Objectives", style="Heading 2")
    add_bullets(doc, [
        "Provide role-specific workspaces for administrators, facilitators and trainees.",
        "Model the organization's real onboarding workflow: a fixed curriculum template (Training Plan) instantiated as an independent, editable copy for every new batch.",
        "Centralize batches, assignments, submissions, sessions, attendance, learning resources, announcements and two distinct kinds of feedback.",
        "Provide operational visibility through metrics, notifications, audit activity and downloadable reports.",
        "Apply authenticated and role-aware access control to sensitive actions and files, at both the frontend route layer and the backend service layer independently.",
    ])
    doc.add_paragraph("Scope and method", style="Heading 2")
    add_bullets(doc, [
        "In scope: React/Vite frontend, Express/TypeScript API, Prisma/PostgreSQL data model, storage abstraction (local/S3), validation (zod), automated tests, and deployment configuration for Vercel/Railway-Render/Neon.",
        "Out of scope: proposed features without code, third-party vendor commitments, production service-level agreements, and business processes not represented in the repository.",
        "Evidence date: repository reviewed 20 July 2026. Backend and frontend production builds both succeeded; the backend automated suite passed 140 tests across 27 files and the frontend suite passed 29 tests across 3 files - all re-run for this revision, not carried forward from an earlier note.",
        "A local PostgreSQL instance has been connected, migrated (10 migrations) and seeded, and real login/JWT plus the real Announcements API have been verified end-to-end against it. A hosted/production database has not yet been proven - see Section 33.",
    ])
    add_gap_callout(doc, "Interpretation", "Where a screen exists without persistence or a backend route, the requirement is labeled Frontend only or Configuration only instead of being described as fully implemented. Where v1.0 recorded something as a gap that is now closed, this document says so explicitly with a green callout.")

    # ---------------- 2. Project Overview ----------------
    page_break(doc)
    add_page_title(doc, "2", "Project Overview and Architecture")
    add_text(doc, "The Trainee Portal is a browser-based training administration system implemented as a TypeScript monorepo. It combines role-specific dashboards with a REST API and a relational data model built around a reusable curriculum-template workflow.")
    add_definition_rows(doc, [
        ("Frontend", "React 18.3, TypeScript, Vite 5, React Router 6, Zustand 5, Tailwind CSS 3. Dev server on :5173, proxies /api to :4000."),
        ("Backend", "Node.js, Express 4, TypeScript REST API, zod validation. Dev server on :4000."),
        ("Persistence", "PostgreSQL 17, accessed through Prisma ORM 5.22; 10 versioned migrations as of the evidence date."),
        ("Files", "Pluggable storage provider interface - local disk (dev default) or Amazon S3 (@aws-sdk/client-s3); Cloudinary is not implemented."),
        ("Authentication", "JWT access token (15-minute default TTL) plus a rotating, hashed refresh token delivered as an HttpOnly cookie."),
        ("Observability", "Structured JSON logger (backend) plus Sentry hook points (@sentry/node backend, @sentry/react frontend) - present in both codebases but inert until a DSN is configured."),
        ("Deployment targets", "Vercel (frontend), Railway preferred or Render (backend), Neon (managed PostgreSQL) - documented and scripted, not yet executed end-to-end in production."),
    ])
    doc.add_paragraph("Logical flow", style="Heading 2")
    add_bullets(doc, [
        "A user signs in, accepts an administrator-issued invitation, or (production only when configured) requests a password-reset email.",
        "The application directs the authenticated user to the dashboard for their assigned role: Admin, Facilitator or Trainee.",
        "An administrator onboards a batch by naming it, selecting one of exactly two Training Plans, and choosing a start date; the plan's full session/assignment/resource/announcement template is copied onto that batch in one database transaction.",
        "The frontend requests protected REST resources; the API validates identity (JWT), role, input shape (zod) and ownership/enrollment before touching the database.",
        "Business records are stored in PostgreSQL via Prisma; uploaded files go through the configured storage provider.",
        "Selected mutations write an AuditLog entry, which is also the source feed for each user's notification panel.",
    ], compact=True)
    add_screenshot_placeholder(doc, "System overview", "application login page beside the three role dashboards, or a simple deployment architecture diagram (Vercel / Railway-Render / Neon)")

    # ---------------- 3. User Roles ----------------
    page_break(doc)
    add_page_title(doc, "3", "User Roles and Access")
    add_definition_rows(doc, [
        ("Admin", "System-wide administrator. Manages users, invitations, batches, Training Plan templates, global assignments, sessions, resources, feedback review, announcements, reports and audit visibility."),
        ("Facilitator", "Training operator. Manages assigned batches, creates assignments/sessions/resources within them, reviews submissions, marks attendance, attaches feedback-form links, and gives trainee performance feedback."),
        ("Trainee", "Learner. Views enrolled batches and assigned work, submits files, joins sessions, uses verified resources, reviews grades, submits eligible feedback forms, and gives facilitator performance feedback."),
    ])
    doc.add_paragraph("Role boundaries", style="Heading 2")
    add_bullets(doc, [
        "The frontend protects Admin, Facilitator and Trainee routes through RequireAuth role checks.",
        "The backend separately protects privileged routes with requireAuth (valid, active JWT subject) and requireRole/requirePermission middleware; ownership and enrollment rules are enforced inside individual services (e.g. assertBatchAccess).",
        "Admin access is generally global. Facilitator modification rights normally require object ownership or batch assignment. Trainee rights normally require enrollment or ownership of the specific record.",
        "The schema contains Role, Permission and RolePermission tables, and a requirePermission() middleware exists and is available, but current route authorization is expressed primarily through the three role names rather than the granular permission set.",
    ])
    add_gap_callout(doc, "Access caveat carried forward from the prior audit", "The v1.0 FRD noted that the assignment-submission roster endpoint was readable by any authenticated user, broader than a least-privilege academic portal would normally require. This was not independently re-verified for v2.0 and should be treated as unconfirmed rather than resolved.")
    add_screenshot_placeholder(doc, "Role dashboards", "one screenshot per role showing the left navigation and dashboard header")

    # ---------------- 4. Information Architecture ----------------
    page_break(doc)
    add_page_title(doc, "4", "Information Architecture and Navigation")
    add_text(doc, "Navigation is standardized per role in frontend/src/constants/navigation.ts. Sessions, Calendar and Events are deliberately merged into a single Sessions & Calendar tab per role, and Discussions has been removed from the Facilitator and Trainee navigation (Admin never had it) - see Section 22.")
    add_definition_rows(doc, [
        ("Admin (/admin)", "Real-time Analytics - Batch Management - Training Plans - Assignments - Sessions & Calendar - Global Resources - Announcements - Feedback - Reports - Audit Logs."),
        ("Facilitator (/facilitator)", "Dashboard - Batches - Assignments - Sessions & Calendar - Resource Library - Announcements - Feedback & Reports - Trainees."),
        ("Trainee (/trainee)", "My Progress - My Batch - Assignments - Sessions & Calendar - Learning Repository - Announcements - My Session Feedback - Facilitators."),
        ("Detail routes", "/admin/training-plans/:planId, /admin/trainees/:name, /facilitator/batches/:batchId, /facilitator/trainees/:name, /assignments/:assignmentId, /{role}/account-settings."),
        ("Contact actions", "Facilitator/trainee contact cards open the operating system's mail client via a mailto: link; there is no in-app messaging."),
    ])
    add_gap_callout(doc, "Deliberate nav-order deviation, accepted", "Training Plans sits third in the Admin sidebar, immediately after Batch Management, rather than grouped with the other role-specific pages at the end. This was a deliberate product decision (it is central to the batch-onboarding workflow) and was flagged, not changed, in the 15 July 2026 audit.")

    # ---------------- 5. Authentication ----------------
    page_break(doc)
    add_page_title(doc, "5", "Authentication and Account Access")
    add_requirements(doc, [
        ("FR-001", "The system shall authenticate an active user by email and password and return the user profile plus a short-lived access token.", "All", "Implemented - API + UI"),
        ("FR-002", "The system shall support a Remember Me choice that extends the refresh-token session lifetime from 7 days to 30 days.", "All", "Implemented - tested"),
        ("FR-003", "The system shall exchange a valid, unexpired refresh-token cookie for a new access token.", "All", "Implemented - tested"),
        ("FR-004", "The system shall revoke the current refresh token and clear the refresh cookie on logout.", "All", "Implemented - API + UI"),
        ("FR-005", "The system shall return the current authenticated user's profile and restore the session during application bootstrap.", "All", "Implemented - API + UI"),
        ("FR-006", "The system shall allow an administrator to create an invitation for an Admin, Facilitator or Trainee email address, valid for 7 days.", "Admin", "Implemented - tested"),
        ("FR-007", "The system shall allow a recipient with a valid, pending, unexpired invite token to set a password and activate the account.", "Invitee", "Implemented - API + UI"),
        ("FR-008", "The system shall accept a forgot-password request for any email address and always respond with the same success message, without revealing whether the account exists.", "Anonymous", "Implemented - tested"),
        ("FR-009", "On a valid forgot-password request, the system shall issue a single-use, SHA-256-hashed reset token valid for one hour and deliver it through the configured email provider.", "Anonymous", "Implemented - tested"),
        ("FR-010", "The system shall set a new password only after verifying an unexpired, unused reset token, then mark the token used.", "Anonymous with token", "Implemented - tested"),
        ("FR-011", "The system shall allow a signed-in user to change their password after verifying their current password.", "All", "Implemented - API + UI"),
        ("FR-012", "New passwords (invite acceptance, reset, change) shall be 8-72 characters and contain at least one letter and one number; login itself does not re-validate this policy against existing passwords.", "All", "Implemented - validated"),
        ("FR-013", "The system shall lock an account for a configurable cooldown period after a configurable number of consecutive failed login attempts, and reset the counter on a successful login.", "System", "Implemented - tested"),
        ("FR-014", "The login screen shall expose demo entry points for the three roles (View as Admin/Facilitator/Trainee) using in-memory session fixtures with no backend calls.", "Demo user", "Implemented - demo only"),
    ], "auth.routes.ts, auth.service.ts, auth.validator.ts, passwordReset.service.ts, authStore.ts, LoginPage.tsx, ResetPasswordPage.tsx, InvitePage.tsx, and auth/password-reset/token-refresh/rate-limit tests.")
    add_resolved_callout(doc, "Resolved since v1.0", "v1.0 recorded an open security gap: Forgot Password accepted an email plus a new password with no verification token. As of 16 July 2026 this has been replaced by the token-verified flow in FR-008 to FR-010 (POST /auth/forgot-password issues the token; POST /auth/reset-password consumes it). Account Settings now calls the current-password-verified change-password endpoint (FR-011).")
    add_gap_callout(doc, "Delivery caveat", "The reset/invite token is delivered through a provider-neutral EmailProvider interface. The only implemented provider logs to the console; no real SES/Postmark/Resend integration exists yet, so in a real deployment the token currently reaches nobody by email unless EXPOSE_AUTH_TOKENS=true is explicitly set (non-production only) or a provider is connected - see Section 33.")
    add_screenshot_placeholder(doc, "Authentication", "login screen, invite acceptance screen, and the forgot-password / reset-password screens")

    # ---------------- 6. User & Profile Management ----------------
    page_break(doc)
    add_page_title(doc, "6", "User and Profile Management")
    add_requirements(doc, [
        ("FR-015", "The system shall allow an administrator to list users with role, active-status, search, sorting and pagination filters.", "Admin", "Implemented - API"),
        ("FR-016", "The system shall allow an authenticated user to view their own profile.", "All", "Implemented - API + UI"),
        ("FR-017", "The system shall allow a user to update their own name, email, phone and location.", "All", "Implemented - API + UI"),
        ("FR-018", "The system shall allow an administrator or the profile owner to retrieve a user record by id.", "Admin, self", "Implemented - API"),
        ("FR-019", "The system shall allow an administrator to change a user's role, active state, company, department and ID number.", "Admin", "Implemented - API"),
        ("FR-020", "The system shall allow an administrator to soft-delete a user, deactivating the account and revoking its outstanding refresh tokens.", "Admin", "Implemented - API"),
        ("FR-021", "The account settings page shall display company/batch metadata as read-only contextual information appropriate to the viewer's role.", "All", "Implemented - UI"),
        ("FR-022", "The user profile data model shall reserve fields for an avatar (storage key, MIME type, size, updated-at), for future use.", "System", "Schema only - no upload route"),
    ], "users.routes.ts, users.service.ts, UserProfile schema, AccountSettingsPage.tsx and the user service/store.")
    doc.add_paragraph("Current limitations", style="Heading 2")
    add_bullets(doc, [
        "No self-registration route is implemented; account creation is invitation-based or seeded for demo/dev use.",
        "Avatar metadata exists in the data model (FR-022), but no complete avatar upload endpoint is exposed in the reviewed routes.",
        "No user-facing administration screen was found for editing the role-permission matrix, even though the underlying Role/Permission/RolePermission tables and a requirePermission() middleware both exist.",
    ])
    add_screenshot_placeholder(doc, "Account settings", "profile information form and role-specific details panel")

    # ---------------- 7. Training Plans ----------------
    page_break(doc)
    add_page_title(doc, "7", "Training Plans and Curriculum Templates", "New module - did not exist in the v1.0 FRD; this is the portal's core business workflow.")
    add_text(doc, "A Training Plan is the organization's fixed curriculum template. Exactly two exist: BA BTech (code ba-btech) and BA MBA (code ba-mba); the program is always BA, and the track is derived from the plan code. Each plan owns its own set of template sessions, template assignments, template resources and template default announcements. Creating a batch against a plan copies that entire template into a new, independent, per-batch set of real records in one database transaction; later edits to a batch's copy never write back to the template, and later edits to the template never propagate to batches already created from it.")
    add_requirements(doc, [
        ("FR-023", "The system shall list all Training Plans (currently exactly two).", "All", "Implemented - API + UI"),
        ("FR-024", "The system shall return a plan's full template detail: header fields plus its complete session, assignment, resource and default-announcement lists.", "All", "Implemented - API + UI"),
        ("FR-025", "The system shall allow an administrator to edit a plan's name, description, duration in months, and its four default minute-of-day timing fields.", "Admin", "Implemented - API + UI"),
        ("FR-026", "The system shall allow an administrator to add, edit or remove a template session (title, agenda, day offset, start/end minute, platform, display order, optional default feedback-form URL).", "Admin", "Implemented - validated"),
        ("FR-027", "A template session's end time shall be later than its start time; the API rejects a session where this does not hold.", "System", "Implemented - validated"),
        ("FR-028", "The system shall allow an administrator to add, edit or remove a template assignment (title, agenda, description, due-day offset, optional related template session).", "Admin", "Implemented - API + UI"),
        ("FR-029", "The system shall allow an administrator to add, edit or remove a template resource (title, category, URL).", "Admin", "Implemented - API + UI"),
        ("FR-030", "The system shall allow an administrator to add, edit or remove a template default announcement (title, message, priority).", "Admin", "Implemented - API + UI"),
        ("FR-031", "Creating a batch against a plan shall instantiate, in one transaction, an independent copy of every template session, assignment, resource and default announcement, with real dates computed from the batch's start date and the plan's day offsets and minute-of-day defaults.", "System", "Implemented - tested"),
        ("FR-032", "Instantiated records shall retain a nullable back-reference to the template row they were generated from, for traceability only; this reference does not create a live link.", "System", "Implemented - schema"),
        ("FR-033", "Editing a Training Plan template shall never modify any batch already generated from it.", "System", "Implemented - tested"),
        ("FR-034", "Editing a batch's copied schedule, assignments, resources or announcements shall never modify the originating Training Plan template.", "System", "Implemented - tested"),
        ("FR-035", "The Admin Training Plan detail page shall display summary statistics (duration, session count, assignment count, number of batches generated) and a template-isolation banner explaining that edits below do not affect existing batches.", "Admin", "Implemented - UI"),
        ("FR-036", "The Admin batch-onboarding flow shall let the administrator deploy (assign) a Training Plan to a new batch directly from the plan's detail page.", "Admin", "Implemented - UI"),
    ], "trainingPlans.routes.ts, trainingPlans.service.ts, trainingPlans.validator.ts, TrainingPlan* schema models, batches.service.ts create(), TrainingPlansPanel.tsx, TrainingPlanDetailPage.tsx, and trainingPlans.service.test.ts / batches.automation.test.ts.")
    add_resolved_callout(doc, "New capability since v1.0", "This entire module - the Training Plan template store, its five CRUD sub-resources, and the batch-instantiation transaction - was implemented after the v1.0 FRD was written (migrations 20260714120000_training_plan_workflow onward) and is documented here for the first time.")
    add_screenshot_placeholder(doc, "Training Plans", "plan list, plan detail with stats and template-isolation banner, add/edit session and assignment dialogs")

    # ---------------- 8. Batch Management ----------------
    page_break(doc)
    add_page_title(doc, "8", "Batch Management")
    add_requirements(doc, [
        ("FR-037", "The system shall list non-deleted batches with search, program/track/status/facilitator/trainee/training-plan filters, sorting and pagination.", "All", "Implemented - API + UI"),
        ("FR-038", "The system shall scope batch lists to a specified facilitator or trainee when the relevant filter is supplied.", "All", "Implemented - tested"),
        ("FR-039", "The system shall allow an administrator to create a batch by supplying a unique code, name, a Training Plan, an optional facilitator (Trainer), an optional start month, and a status.", "Admin", "Implemented - tested"),
        ("FR-040", "The batch's program and track shall be derived server-side from the selected Training Plan, not accepted from the client.", "System", "Implemented - service"),
        ("FR-041", "The batch's end date shall always be server-computed at creation time from the generated schedule (the date of the last scheduled session, falling back to duration-in-months times 30 days); it is never accepted as create input.", "System", "Implemented - service"),
        ("FR-042", "A batch's Training Plan link shall be immutable after creation; the end date becomes independently editable afterward so an administrator can shorten or extend one batch's schedule without touching the template.", "Admin", "Implemented - validated"),
        ("FR-043", "The system shall allow an administrator to update or soft-delete (archive) a batch.", "Admin", "Implemented - tested"),
        ("FR-044", "The system shall calculate batch metrics: average score, completion rate, submission rate, attendance rate and average feedback rating.", "Admin, Facilitator", "Implemented - API"),
        ("FR-045", "The system shall allow Admin or Facilitator users to list trainees in an accessible batch.", "Admin, Facilitator", "Implemented - scoped"),
        ("FR-046", "The system shall allow Admin or Facilitator users to enroll or remove a trainee from a batch without deleting historical enrollment rows (removal sets removedAt; a later enrollment reactivates the same row).", "Admin, Facilitator", "Implemented - API"),
        ("FR-047", "The system shall provide per-trainee batch statistics - assignment progress, average grade, attendance percentage, latest submission and feedback status - restricted to admin or the batch's own facilitator.", "Admin, owner Facilitator", "Implemented - tested"),
    ], "batches.routes.ts, batches.service.ts, Batch and BatchTrainee schema, batches.crud/automation/facilitator-scope/trainee-scope tests.")
    add_resolved_callout(doc, "Changed since v1.0", "v1.0 documented batch creation as accepting a raw program and track directly. The current workflow (FR-039 to FR-041) instead requires selecting a Training Plan, from which program, track and end date are all derived - the administrator now only supplies a name, a plan and a start date, matching the organization's actual onboarding process.")
    add_screenshot_placeholder(doc, "Batch management", "Admin batch table, the create-batch modal (name + plan + start date), and the Facilitator batch-detail performance page")

    # ---------------- 9. Assignment Management ----------------
    page_break(doc)
    add_page_title(doc, "9", "Assignment Management")
    add_requirements(doc, [
        ("FR-048", "The system shall list non-deleted assignments with batch/status filters, search, sorting and pagination.", "All", "Implemented - API + UI"),
        ("FR-049", "The system shall allow Admin or Facilitator users to create an assignment for one or more batches in a single call, with title, agenda, description, deadline, status, optional related session, and an optional instructions file.", "Admin, Facilitator", "Implemented - tested"),
        ("FR-050", "The first selected batch shall be retained as the assignment's primary batch for backward-compatible single-batch reads; the authoritative multi-batch set lives in a join table.", "System", "Implemented - service/schema"),
        ("FR-051", "The creator shall be recorded as the assignment's facilitator when created by a facilitator; assignments generated from a Training Plan carry no facilitator by default.", "System", "Implemented - service/schema"),
        ("FR-052", "The system shall allow an administrator or owning facilitator to update assignment data, its batch links, its related session, or its instructions file.", "Admin, owner Facilitator", "Implemented - authorized"),
        ("FR-053", "The system shall allow an administrator or owning facilitator to soft-delete an assignment.", "Admin, owner Facilitator", "Implemented - authorized"),
        ("FR-054", "The system shall allow an administrator, an owning or co-facilitator (any facilitator on one of its batches), or an enrolled trainee to view or download the instruction attachment.", "Scoped users", "Implemented - tested"),
        ("FR-055", "The submission roster for an assignment shall return exactly one row per enrolled trainee, including a synthetic Not Started placeholder row (id: null) for a trainee who has not yet submitted.", "Authenticated", "Implemented - API"),
        ("FR-056", "An assignment may optionally link to a Session via sessionId, surfaced as Related Session in the assignment view and reciprocally as the session's related assignment.", "Admin, Facilitator", "Implemented - schema/UI"),
        ("FR-057", "The assignment dashboard shall support selecting multiple assignments and issuing a bulk deadline extension as individually authorized API calls.", "Admin, Facilitator", "Implemented - UI/API calls"),
    ], "assignments.routes.ts, assignments.service.ts, assignments.validator.ts, Assignment/AssignmentBatch schema, and assignments.multi-batch/response-shape tests.")
    add_screenshot_placeholder(doc, "Assignment management", "assignment table, multi-batch selection, related-session field, and the create-assignment dialog")

    # ---------------- 10. Submission & Grading ----------------
    page_break(doc)
    add_page_title(doc, "10", "Submission and Grading")
    add_requirements(doc, [
        ("FR-058", "The system shall allow an actively enrolled trainee to submit their own work for an assigned assignment.", "Trainee", "Implemented - authorized"),
        ("FR-059", "The system shall mark a new submission Under Review or Late according to the assignment deadline at the moment of submission.", "System", "Implemented - service"),
        ("FR-060", "There shall be at most one submission per trainee and assignment; resubmission upserts the same row.", "System", "Implemented - schema/tested"),
        ("FR-061", "The system shall allow the trainee to attach a file to their own submission.", "Trainee", "Implemented - upload API"),
        ("FR-062", "Replacement of a current attachment shall be allowed only before the deadline; the superseded attachment is retained (isCurrent set to false) rather than deleted, so submission history is preserved.", "Trainee", "Implemented - tested"),
        ("FR-063", "The system shall reject replacement of an existing attachment after the deadline, while still allowing a first late submission.", "Trainee", "Implemented - tested"),
        ("FR-064", "The system shall allow an administrator or owning facilitator to enter a grade from 0 to 100 (two decimal places) and optional written feedback.", "Admin, owner Facilitator", "Implemented - API + UI"),
        ("FR-065", "Submission retrieval and attachment download shall be restricted to the trainee owner, the assignment's facilitator, or an administrator.", "Scoped users", "Implemented - tested"),
    ], "submissions.routes.ts, submissions.service.ts, submissions.validator.ts, Submission/SubmissionAttachment schema, AssignmentDetailPage.tsx, and submissions.authorization/resubmission tests.")
    add_gap_callout(doc, "UI placeholder retained from v1.0", "The Facilitator dashboard still includes a text-only submission-summary download helper labeled as a placeholder; the actual uploaded submission files are handled entirely through the secured attachment routes in FR-061 to FR-065.")

    # ---------------- 11. Sessions and Calendar ----------------
    page_break(doc)
    add_page_title(doc, "11", "Sessions and Calendar")
    add_requirements(doc, [
        ("FR-066", "The system shall list sessions with batch/status filters, search, sorting and pagination.", "All", "Implemented - API + UI"),
        ("FR-067", "The system shall allow Admin or Facilitator users to schedule a batch session with title, agenda, date/time, duration in minutes (default 120), platform, optional meeting URL and status.", "Admin, Facilitator", "Implemented - API + UI"),
        ("FR-068", "The Trainer (facilitator) on a session shall be optional, matching the organization's Training-Plan-generated workflow where most sessions carry no pre-assigned trainer.", "System", "Implemented - schema"),
        ("FR-069", "The system shall allow an administrator or owning facilitator to update, reschedule, cancel or soft-delete a session.", "Admin, owner Facilitator", "Implemented - API + UI"),
        ("FR-070", "The system shall support Google Meet, Microsoft Teams, Zoom and Other as session platforms.", "Admin, Facilitator", "Implemented - enum/validation"),
        ("FR-071", "The system shall expose session states Upcoming, Live, Completed, Cancelled and Rescheduled.", "All", "Implemented - enum/validation"),
        ("FR-072", "The system shall provide a combined calendar feed of sessions and assignment deadlines, filterable by batch and by event type.", "All", "Implemented - tested"),
        ("FR-073", "Calendar results shall be scoped to all batches for Admin, managed batches for Facilitator, and active enrollments for Trainee.", "All", "Implemented - tested"),
        ("FR-074", "A trainee shall be able to open a valid meeting link from an eligible session.", "Trainee", "Implemented - UI"),
    ], "sessions.routes.ts, calendar.routes.ts, sessions.service.ts, calendar.service.ts, Session schema, SessionsCalendarView.tsx and calendar tests.")
    add_screenshot_placeholder(doc, "Sessions and calendar", "list/calendar toggle showing session timing, related assignment, and feedback-form link")

    # ---------------- 12. Attendance ----------------
    page_break(doc)
    add_page_title(doc, "12", "Attendance Management")
    add_requirements(doc, [
        ("FR-075", "The system shall allow Admin or Facilitator users to retrieve the attendance roster for a session.", "Admin, Facilitator", "Implemented - API"),
        ("FR-076", "The system shall allow Admin or Facilitator users to bulk-mark attendance for one or more trainees in a session in a single call.", "Admin, Facilitator", "Implemented - API"),
        ("FR-077", "The system shall allow an administrator or owning facilitator to correct a single attendance record after the fact.", "Admin, owner Facilitator", "Implemented - API"),
        ("FR-078", "Attendance status shall be Present, Absent, Late or Excused.", "System", "Implemented - enum/validation"),
        ("FR-079", "Attendance shall be unique per session and trainee, and shall record who marked it and when.", "System", "Implemented - schema"),
        ("FR-080", "Batch and per-trainee analytics shall treat Present and Late records as attended when calculating attendance rate.", "System", "Implemented - service/test"),
    ], "attendance.routes.ts, sessions.routes.ts attendance sub-routes, attendance.service.ts, Attendance schema, batch metrics.")
    add_screenshot_placeholder(doc, "Attendance", "session attendance roster with status controls and summary counts")

    # ---------------- 13. Resource Library ----------------
    page_break(doc)
    add_page_title(doc, "13", "Resource Library")
    add_requirements(doc, [
        ("FR-081", "The system shall list resources with batch, category, verified-status, search, sorting and pagination filters.", "All", "Implemented - API + UI"),
        ("FR-082", "The system shall allow Admin or Facilitator users to upload a file with title, category, version (default v1.0) and optional batch scope.", "Admin, Facilitator", "Implemented - upload API"),
        ("FR-083", "A resource shall carry either an uploaded file (storage key, MIME type, byte size) or an external URL (for example, one copied from a Training Plan template), with exactly one populated - enforced at the application layer, not the database.", "System", "Implemented - schema/service"),
        ("FR-084", "The system shall record the uploader and increment a download count on every successful download.", "System", "Implemented - schema/service"),
        ("FR-085", "The uploader or an administrator shall be able to update resource metadata or soft-delete a resource.", "Admin, owner uploader", "Implemented - authorized"),
        ("FR-086", "An administrator shall be able to set the verification status that gates trainee-facing visibility in the UI.", "Admin", "Implemented - update API/UI"),
        ("FR-087", "The system shall stream an available resource file on download and increment its download count.", "All", "Implemented - API"),
        ("FR-088", "A resource instantiated from a Training Plan template shall retain a nullable back-reference to that template resource, for traceability only.", "System", "Implemented - schema"),
        ("FR-089", "The Admin interface shall support bulk verification and bulk deletion, each executed as individually secured API calls.", "Admin", "Implemented - UI orchestration"),
    ], "resources.routes.ts, resources.service.ts, resources.validator.ts, Resource/TrainingPlanResource schema, storage providers.")
    add_gap_callout(doc, "Visibility rule caveat carried forward from the prior audit", "The UI states that resources become visible to trainees once verified, but the generic list API does not itself force verified=true by caller role - enforcement depends on the frontend query. This was not independently re-verified for v2.0.")
    add_screenshot_placeholder(doc, "Resource library", "resource filters, verified badge, upload dialog and download action")

    # ---------------- 14. Session Feedback ----------------
    page_break(doc)
    add_page_title(doc, "14", "Session Feedback (Form Links)", "New module - not documented as a distinct feature in v1.0.")
    add_text(doc, "Session Feedback is an external form link (typically Google Forms or Microsoft Forms) attached to a Session, with per-user submission tracking used for status and simple completion stats. It is unrelated to the facilitator/trainee performance-rating feedback in Section 16 - the portal never collects the external form's actual answers, only who has clicked through.")
    add_requirements(doc, [
        ("FR-090", "The system shall allow an administrator or owning facilitator to attach one feedback-form link to a session: a name, description, a URL validated as a well-formed URL, and an audience.", "Admin, owner Facilitator", "Implemented - tested"),
        ("FR-091", "The audience shall be Trainees, Facilitators or Both, gating who can see and submit the form; the admin and the owning facilitator can always see it regardless of audience.", "System", "Implemented - tested"),
        ("FR-092", "The system shall allow an administrator or owning facilitator to edit or remove a session's feedback form.", "Admin, owner Facilitator", "Implemented - API + UI"),
        ("FR-093", "The system shall allow an administrator, the owning facilitator, or an enrolled trainee to retrieve the form's link plus submission stats (submitted count and whether the caller has submitted).", "Scoped users", "Implemented - tested"),
        ("FR-094", "An eligible trainee or facilitator shall be able to mark their own submission as complete; the action is idempotent (unique per form and submitter) and does not collect the form's actual responses.", "Trainee, Facilitator", "Implemented - tested"),
        ("FR-095", "A session's default feedback-form URL may be pre-populated from its originating Training Plan template session and copied forward when a batch is created.", "System", "Implemented - schema/service"),
    ], "sessionFeedback.routes.ts, sessionFeedback.service.ts, sessionFeedback.validator.ts, SessionFeedbackForm/SessionFeedbackSubmission schema, SessionFeedbackCell.tsx, sessionFeedback.service.test.ts.")
    add_screenshot_placeholder(doc, "Session feedback", "session row with attach/edit/copy/open feedback-form controls and submitted-count")

    # ---------------- 15. Assignment Feedback ----------------
    page_break(doc)
    add_page_title(doc, "15", "Assignment Feedback (Form Links)", "New module, added 16 July 2026 - resolves a conflict flagged in the prior audit.")
    add_text(doc, "Assignment Feedback mirrors the Session Feedback pattern above, but attaches directly to an Assignment rather than to a Session - covering feedback that belongs to a specific piece of work (an assignment with no related session, or a work-specific survey) independently of any session-level form.")
    add_requirements(doc, [
        ("FR-096", "The system shall allow an administrator or owning facilitator to attach a feedback-form link directly to an assignment, independent of any related session's form.", "Admin, owner Facilitator", "Implemented - tested"),
        ("FR-097", "Assignment Feedback shall use the same audience gating (Trainees/Facilitators/Both) and per-user idempotent submission-tracking model as Session Feedback.", "System", "Implemented - tested"),
        ("FR-098", "The system shall allow an administrator or owning facilitator to edit or remove an assignment's feedback form.", "Admin, owner Facilitator", "Implemented - API + UI"),
        ("FR-099", "An assignment that also has a related session shall track its own feedback form and that session's feedback form as fully independent records.", "System", "Implemented - schema"),
    ], "assignmentFeedback.routes.ts, assignmentFeedback.service.ts, AssignmentFeedbackForm/AssignmentFeedbackSubmission schema, AssignmentDetailPage.tsx, assignmentFeedback.service.test.ts.")
    add_resolved_callout(doc, "Resolved conflict from the 15 July 2026 audit", "That audit recorded item D6b as a CONFLICT: feedback-form links attached only to sessions, with no direct assignment-level link, pending a user decision. The decision was made and implemented the next day - this entire module is the result.")

    # ---------------- 16. Performance Feedback ----------------
    page_break(doc)
    add_page_title(doc, "16", "Facilitator-Trainee Performance Feedback")
    add_text(doc, "This is the rating-based feedback module (FeedbackEntry) - distinct from the two form-link features above. It captures a category, an integer rating and an optional comment, directionally between a facilitator and a trainee in a shared batch.")
    add_requirements(doc, [
        ("FR-100", "The system shall list feedback entries with batch, trainee, facilitator and direction filters, sortable by rating or date.", "Scoped users", "Implemented - API + UI"),
        ("FR-101", "The system shall allow Admin or Facilitator users to submit rating-based feedback about an actively enrolled trainee.", "Admin, Facilitator", "Implemented - tested"),
        ("FR-102", "The system shall allow a trainee to submit feedback only about the facilitator of an actively enrolled batch.", "Trainee", "Implemented - tested"),
        ("FR-103", "Feedback shall include a category, an integer rating from 1 to 5, and an optional comment.", "All", "Implemented - validated"),
        ("FR-104", "The system shall stamp feedback direction automatically as Facilitator-to-Trainee or Trainee-to-Facilitator based on who submitted it.", "System", "Implemented - tested"),
        ("FR-105", "Feedback records shall be append-only; no update or delete endpoint is exposed.", "All", "Implemented - route design"),
        ("FR-106", "A trainee's feedback list shall be restricted to that trainee even if a different trainee filter is requested.", "Trainee", "Implemented - tested"),
    ], "feedback.routes.ts, feedback.service.ts, feedback.validator.ts, FeedbackEntry schema, feedback.direction.test.ts.")
    add_screenshot_placeholder(doc, "Performance feedback", "Facilitator feedback form and Trainee feedback/grades panel")

    # ---------------- 17. Announcements ----------------
    page_break(doc)
    add_page_title(doc, "17", "Announcements", "Reclassified from Frontend only (v1.0) to fully implemented and database-backed.")
    add_requirements(doc, [
        ("FR-107", "The system shall persist announcements with title, message, priority (Normal/Important/Critical), a free-text audience label, an optional batch scope (null meaning global), a pinned flag, and optional schedule/expiry timestamps.", "System", "Implemented - schema/API"),
        ("FR-108", "The system shall allow an administrator to post an announcement to any scope, and a facilitator to post only to their own batches.", "Admin, Facilitator", "Implemented - tested"),
        ("FR-109", "The system shall list announcements visible to the caller: global announcements plus the caller's own batches for facilitators and trainees; an administrator sees all announcements.", "All", "Implemented - API + UI"),
        ("FR-110", "The system shall allow the author or an administrator to edit or soft-delete an announcement.", "Admin, author Facilitator", "Implemented - API"),
        ("FR-111", "The system shall allow a user to mark an announcement as read; read state is stored per user, independently of the announcement's content.", "All", "Implemented - API + UI"),
        ("FR-112", "An announcement instantiated from a Training Plan's default announcements at batch creation shall retain a nullable back-reference to that template announcement.", "System", "Implemented - schema"),
        ("FR-113", "Demo Mode shall maintain full parity with the real API: identical fixture data and scoped handlers so the announcements screen behaves the same whether backed by Demo Mode or PostgreSQL.", "All", "Implemented - tested"),
    ], "announcements.routes.ts, announcements.service.ts, announcements.validator.ts, Announcement/AnnouncementRead schema, announcementsStore.ts, announcements.service.test.ts, and demoMode.ts scoped handlers.")
    add_resolved_callout(doc, "Resolved since v1.0", "v1.0 recorded Announcements as Frontend only: 'No /api/announcements routes exist. Announcement posts and reads are not database-backed.' As of the announcements-API work (verified against a live local PostgreSQL instance on 17 July 2026 - real role/batch scoping, real per-user readByCount increment), this is fully implemented and is the module the prior FRD called out as the last remaining mock-data holdout. It is not a holdout anymore.")
    add_screenshot_placeholder(doc, "Announcements", "announcement composer with priority/audience/pin controls, and the trainee announcements feed")

    # ---------------- 18. Notifications & Audit ----------------
    page_break(doc)
    add_page_title(doc, "18", "Notifications and Audit Logs")
    add_requirements(doc, [
        ("FR-114", "The system shall record supported business events to an audit log with actor, event type, message, module, previous/new value snapshots, IP address and timestamp.", "System", "Implemented - service"),
        ("FR-115", "The Admin dashboard shall list and filter audit activity by message, user, module and event attributes.", "Admin", "Implemented - UI/API data"),
        ("FR-116", "The system shall derive a user's notification feed from audit events plus per-user read state.", "All", "Implemented - API"),
        ("FR-117", "The notification feed shall return an unread count with paginated notification items, optionally filtered to unread only.", "All", "Implemented - API + UI"),
        ("FR-118", "A user shall be able to mark one notification as read.", "All", "Implemented - API + UI"),
        ("FR-119", "A user shall be able to mark all visible notifications as read in one call.", "All", "Implemented - API + UI"),
        ("FR-120", "Notification read state shall be stored separately for each user and audit event, and shall never mutate the underlying audit log entry.", "System", "Implemented - schema"),
    ], "audit.ts, notifications.routes.ts, notifications.service.ts, AuditLog/NotificationRead schema, NotificationPanel.tsx, audit-coverage.test.ts.")
    add_gap_callout(doc, "Coverage boundary", "Audit logging covers users, batches, assignments, submissions, resources, sessions and performance feedback. Announcements now persist to the database (Section 17) but were not confirmed in this pass to write their own audit-log entries; discussions have no backend at all (Section 22) and therefore generate no audit events.")
    add_screenshot_placeholder(doc, "Notifications and audit", "notification panel with unread states and Admin audit-log filters")

    # ---------------- 19. Dashboards ----------------
    page_break(doc)
    add_page_title(doc, "19", "Dashboards and Analytics")
    add_requirements(doc, [
        ("FR-121", "The Admin dashboard shall present system analytics, batch performance comparison, upcoming deadlines, recent activity and quick actions.", "Admin", "Implemented - UI"),
        ("FR-122", "The Facilitator dashboard shall present assigned batches, trainee counts, score/completion/attendance metrics, pending reviews and recent submissions.", "Facilitator", "Implemented - UI/API data"),
        ("FR-123", "The Facilitator batch-detail page shall present trainee, assignment, attendance, session and submission summaries for one batch.", "Facilitator", "Implemented - UI/API data"),
        ("FR-124", "The Trainee dashboard shall present personal progress, assignments, sessions, resources, grades/feedback and facilitator contacts.", "Trainee", "Implemented - UI/API data"),
        ("FR-125", "Dashboard metrics shall display neutral/empty states when required data is unavailable rather than inventing values.", "All", "Implemented - UI patterns"),
        ("FR-126", "The application shall provide stat cards, progress bars, status badges and simple bar charts for summarized operational data.", "All", "Implemented - components"),
    ], "role dashboard pages, FacilitatorBatchDetailPage.tsx, StatCard/ProgressBar/BarChart/EmptyState components, batch metrics service.")
    add_gap_callout(doc, "Meaning of real-time", "The Admin navigation uses the label Real-time Analytics, but the repository has no WebSocket or push layer. Values refresh through ordinary API/state updates, not a live channel.")
    add_screenshot_placeholder(doc, "Analytics", "Admin analytics dashboard and Facilitator/Trainee progress cards")

    # ---------------- 20. Reports ----------------
    page_break(doc)
    add_page_title(doc, "20", "Reports")
    add_requirements(doc, [
        ("FR-127", "The Admin interface shall generate Attendance, Assignment, Performance, Feedback, Session, Resource Usage and Audit report datasets from live API data.", "Admin", "Implemented - client"),
        ("FR-128", "The Admin interface shall filter applicable report datasets by date range and export them as CSV files.", "Admin", "Implemented - client"),
        ("FR-129", "The Admin interface shall open a print-ready report view that can be saved as PDF through the browser's print dialog.", "Admin", "Implemented - client"),
        ("FR-130", "The Admin interface shall allow Daily, Weekly or Monthly report schedules to be added and removed as local configuration entries.", "Admin", "Configuration only"),
    ], "AdminDashboardPage.tsx report builder/export/schedule functions.")
    add_gap_callout(doc, "Still open: scheduled reports do not run", "FR-130's schedules are stored as configuration only - no backend scheduler or job runner exists to actually generate and send a report on a Daily/Weekly/Monthly cadence. This gap is unchanged from v1.0.")

    # ---------------- 21. Common UX ----------------
    page_break(doc)
    add_page_title(doc, "21", "Common User Experience Functions")
    add_requirements(doc, [
        ("FR-131", "The application shall route authenticated users to role-specific dashboards and redirect unauthorized role access.", "All", "Implemented - UI/API"),
        ("FR-132", "The dashboards shall provide consistent role-specific sidebar navigation and page headings.", "All", "Implemented - UI"),
        ("FR-133", "List screens shall provide relevant search, filter, sort, pagination and empty-state behavior.", "All", "Implemented - UI/API"),
        ("FR-134", "The UI shall use confirmation dialogs for destructive actions and saving indicators for asynchronous forms.", "All", "Implemented - UI"),
        ("FR-135", "The UI shall display success/error toast messages for completed or failed actions.", "All", "Implemented - UI"),
        ("FR-136", "The Admin interface shall provide global search (Ctrl+K) across available client-side entities and navigation results.", "Admin", "Implemented - UI"),
        ("FR-137", "The application shall lazy-load major dashboard routes and show a loading fallback during route retrieval.", "All", "Implemented - App.tsx"),
        ("FR-138", "The application shall show a recoverable error fallback (with a Reload page action) when an uncaught frontend render error occurs, and shall report that error through the Sentry hook point when configured.", "All", "Implemented - ErrorBoundary"),
        ("FR-139", "Every modal dialog shall support Escape-to-close and click-outside-to-close.", "All", "Implemented - useEscapeKey"),
    ], "App.tsx, RequireAuth, navigation constants, reusable components, dashboard list screens, Toast, ErrorBoundary.tsx.")
    add_gap_callout(doc, "Contact/reminder behavior", "Several Contact or Reminder actions use mailto: links or local audit/toast behavior; there is no internal messaging, SMS or email delivery workflow beyond the auth-related emails in Section 5.")

    # ---------------- 22. Discussions ----------------
    page_break(doc)
    add_page_title(doc, "22", "Discussions (Dormant Schema)")
    add_text(doc, "DiscussionThread and DiscussionMessage tables exist in the Prisma schema - a thread per batch, authored messages with a snapshotted author role - but there are no backend routes for either model and no frontend UI references them. The Discussions navigation item has been intentionally removed from the Facilitator and Trainee sidebars (Admin never had one); a direct visit to a legacy discussions URL falls through to the catch-all route.")
    add_requirements(doc, [
        ("FR-140", "The data model shall retain DiscussionThread and DiscussionMessage tables for potential future use.", "System", "Schema only - dormant"),
        ("FR-141", "No Discussions capability shall be exposed in any role's navigation.", "System", "Implemented - by design"),
    ], "schema.prisma DiscussionThread/DiscussionMessage models; navigation.ts (no Discussions entry); App.tsx catch-all route.")
    add_gap_callout(doc, "Not a bug", "This is a deliberate, documented product decision, not an oversight - the schema was left in place in case discussions are revisited later, but the feature itself was intentionally removed from scope.")

    # ---------------- 23. BR Access & Identity ----------------
    page_break(doc)
    add_page_title(doc, "23", "Business Rules - Access and Identity")
    add_rules(doc, [
        ("BR-001", "Only active, non-deleted accounts may authenticate or refresh a session.", "requireAuth middleware, auth.service.ts"),
        ("BR-002", "Login is rate-limited to 10 attempts per 15 minutes per client, counting only failed attempts; a separate per-account lockout (default 5 consecutive failures, 15-minute cooldown, both env-configurable) applies independently.", "middleware/rateLimit.ts, auth.service.ts login()"),
        ("BR-003", "Forgot-password requests are rate-limited to 5 per hour per client and always return the same success response regardless of whether the email exists.", "middleware/rateLimit.ts, auth.service.ts"),
        ("BR-004", "New passwords must be 8-72 characters and contain at least one letter and one number; 72 is the ceiling because bcrypt silently truncates beyond it.", "auth.validator.ts passwordPolicy"),
        ("BR-005", "A password reset token is single-use, hashed with SHA-256 at rest, and expires exactly one hour after issue.", "PasswordResetToken schema, passwordReset.service.ts"),
        ("BR-006", "An invitation token is valid only while Pending and before its 7-day expiry; only an administrator may create one.", "UserInvite schema, auth.service.ts, auth.routes.ts"),
        ("BR-007", "Refresh, invite and password-reset tokens are stored only as a SHA-256 hash; the raw token is never persisted.", "RefreshToken/UserInvite/PasswordResetToken schema"),
        ("BR-008", "Access tokens expire in 15 minutes by default; refresh tokens expire in 7 days by default, or 30 days when Remember Me is selected.", "config/env.ts JWT_ACCESS_EXPIRES_IN, REFRESH_TOKEN_TTL_DAYS, REFRESH_TOKEN_REMEMBER_TTL_DAYS"),
        ("BR-009", "In production, invite and reset tokens are never echoed back in an API response; they are reachable only through the configured email provider (or a deliberately set EXPOSE_AUTH_TOKENS=true outside production).", "config/env.ts exposeAuthTokens, invite-token-exposure.test.ts"),
        ("BR-010", "Access-token-protected routes require a Bearer token and re-check on every request that the corresponding database user still exists, is active, and is not soft-deleted.", "requireAuth middleware"),
        ("BR-011", "Frontend route checks do not replace backend authorization; both layers are present and independently enforced.", "App.tsx RequireAuth, backend route middleware"),
        ("BR-012", "Admin may act globally. Facilitator modification rights generally require object ownership or batch assignment. Trainee rights generally require ownership or active enrollment.", "Service-layer authorization across modules"),
        ("BR-013", "Deleting users, batches, assignments, sessions, resources or announcements is implemented as soft deletion (deletedAt); records remain in storage for history.", "Prisma schema / services"),
        ("BR-014", "Deleting (soft) a user also deactivates the account and revokes its active refresh tokens.", "users.service.ts"),
    ])

    # ---------------- 24. BR Training Plan / Curriculum ----------------
    page_break(doc)
    add_page_title(doc, "24", "Business Rules - Training Plan and Curriculum Generation")
    add_rules(doc, [
        ("BR-015", "Exactly two Training Plans exist: ba-btech and ba-mba. Program is always BA; track is derived from the plan code.", "TrainingPlan schema, batches.service.ts deriveProgramTrack"),
        ("BR-016", "Working-day scheduling (nthWorkingDay) counts Monday-Friday only; if a batch's start date falls on a weekend, the schedule rolls forward to the next Monday. The identical algorithm is implemented once in the backend and mirrored in the frontend's Demo Mode fixture generator.", "batches.service.ts, frontend demoData.ts"),
        ("BR-017", "A template session's day offset and minute-of-day fields (defaults: sessions run 14:30-16:30, assignment work opens 09:30, deadline 23:59) are combined with the batch's start date to compute every instantiated Session's scheduledAt and every instantiated Assignment's deadline in local wall-clock time, not UTC.", "TrainingPlan default* fields, batches.service.ts create()"),
        ("BR-018", "Each standard plan spans approximately two months and generates 42 sessions, one per working day.", "seed.ts, trainingPlans fixtures"),
        ("BR-019", "Near-daily case-study assignments are generated - approximately 36 per plan (42 working days minus orientation/wrap-up days, matched by a regex over the session title) - rather than a small fixed count.", "batches.service.ts assignment generation"),
        ("BR-020", "A batch's end date is the date of its last scheduled session; if no sessions exist, it falls back to the Training Plan's duration in months times 30 days.", "batches.service.ts create()"),
        ("BR-021", "A Training Plan's trainingPlanId is immutable on a batch after creation - the instantiation transaction runs exactly once, at create time.", "batches.validator.ts updateBatchSchema"),
        ("BR-022", "Template edits never propagate to previously generated batches, and batch-level edits never write back to the template.", "Separate service layers; no reverse relation write path"),
        ("BR-023", "Trainer/facilitator assignment is optional at both the batch level and the session level; the acting administrator's id is used only for audit-ownership fields, never inferred as a Trainer.", "Batch.facilitatorId, Session.facilitatorId (both nullable)"),
    ])

    # ---------------- 25. BR Training Operations ----------------
    page_break(doc)
    add_page_title(doc, "25", "Business Rules - Training Operations")
    add_rules(doc, [
        ("BR-024", "A trainee may be enrolled once per batch; removal is recorded with removedAt rather than deleting the row, and a later enrollment reactivates it.", "BatchTrainee schema/service"),
        ("BR-025", "A multi-batch assignment must reference at least one batch and stores one primary batch for backward compatibility with single-batch reads.", "assignments.validator.ts batchIdsField"),
        ("BR-026", "There is at most one submission per assignment and trainee (unique constraint on assignmentId + traineeId).", "Submission schema"),
        ("BR-027", "Submission status is determined by timing and review state: Not Started, Under Review, Completed or Late.", "Submission enum/service"),
        ("BR-028", "Only the current attachment is served by the default download route; prior replacement files remain stored as superseded history (isCurrent = false), never deleted.", "SubmissionAttachment schema/service"),
        ("BR-029", "Grades, when supplied, must be between 0 and 100, stored with two decimal places.", "submissions.validator.ts gradeSubmissionSchema"),
        ("BR-030", "Attendance is unique by session and trainee; status is Present, Absent, Late or Excused, and Present/Late both count as attended for rate calculations.", "Attendance schema/validator, n-plus-one/batch metrics tests"),
        ("BR-031", "A resource carries either an uploaded file or an external URL, never both - enforced in the service layer, not by a database constraint.", "resources.service.ts"),
        ("BR-032", "Resource update/delete requires uploader ownership unless the actor is an administrator; every successful download increments the resource's download counter.", "resources.service.ts"),
        ("BR-033", "Calendar and batch-roster access follow the same tiering everywhere: Admin sees globally, Facilitator sees managed batches, Trainee sees enrolled batches.", "calendar.service.ts, batches.service.ts assertBatchAccess"),
        ("BR-034", "Audit-derived notification read state is stored per user and never alters the underlying audit event.", "NotificationRead schema/service"),
    ])

    # ---------------- 26. BR Feedback & Communication ----------------
    page_break(doc)
    add_page_title(doc, "26", "Business Rules - Feedback and Communication")
    add_rules(doc, [
        ("BR-035", "Performance-feedback rating must be an integer from 1 to 5; feedback cannot be edited or deleted through the API once submitted.", "feedback.validator.ts, feedback.routes.ts"),
        ("BR-036", "Session Feedback and Assignment Feedback both use the same three-value audience (Trainees, Facilitators, Both) to gate visibility and submission eligibility; the admin and the owning facilitator can always see the form regardless of audience.", "sessionFeedback/assignmentFeedback isRespondentFor logic"),
        ("BR-037", "A feedback-form URL (session-level, assignment-level, or a Training Plan template default) must be a well-formed URL; the backend rejects anything else with zod's .url() check.", "sessionFeedback.validator.ts, assignmentFeedback validator, trainingPlans.validator.ts"),
        ("BR-038", "Marking a session or assignment feedback form as submitted is idempotent per submitter and does not capture the external form's actual answers - only click-through/completion state.", "SessionFeedbackSubmission / AssignmentFeedbackSubmission unique constraint"),
        ("BR-039", "An announcement is either global (batchId null) or scoped to one batch; a facilitator may only author batch-scoped announcements for their own batches, never global ones.", "announcements.service.ts, announcements.routes.ts requireRole"),
        ("BR-040", "An announcement's read state is recorded per user per announcement and does not affect the announcement's own edit history.", "AnnouncementRead schema"),
    ])

    # ---------------- 27. NFR Security ----------------
    page_break(doc)
    add_page_title(doc, "27", "Non-Functional Requirements - Security")
    add_requirements(doc, [
        ("NFR-001", "All business API routes except documented public endpoints (login, refresh, invite-accept, forgot-password, reset-password, health) shall require authentication.", "Security", "Implemented"),
        ("NFR-002", "Privileged API operations shall enforce role and ownership/enrollment checks on the server, independent of frontend route guards.", "Security", "Implemented; one caveat carried forward (Section 3)"),
        ("NFR-003", "Passwords shall be stored using bcrypt hashing (default 12 salt rounds, configurable 8-15) and shall never be returned in API DTOs.", "Security", "Implemented"),
        ("NFR-004", "Refresh, invite and password-reset tokens shall be stored as SHA-256 hashes; refresh tokens support rotation-on-refresh and explicit revocation on logout or account deletion.", "Security", "Implemented - tested"),
        ("NFR-005", "Production cookies shall be HttpOnly and Secure, with sameSite set to a value that keeps the refresh cookie working across a cross-domain frontend/backend deployment.", "Security", "Implemented by config"),
        ("NFR-006", "The API shall use Helmet, explicit credentialed CORS configuration, gzip response compression, request size limits, and three tiers of rate limiting (general API 300/15min, login 10/15min counting only failures, forgot-password 5/hour).", "Security", "Implemented - tested"),
        ("NFR-007", "Uploaded files shall be rejected by a denylist of executable/script extensions rather than a narrow allowlist (submissions legitimately include arbitrary code, archives, docs and media); size is capped by MAX_UPLOAD_SIZE_MB (default 10 MB); local-disk filenames are always randomly generated, never derived from user input.", "Security", "Implemented"),
        ("NFR-008", "Production error responses shall not expose stack traces, connection strings or dependency internals; a database-layer failure returns a generic 503 with no infrastructure detail.", "Privacy", "Implemented - tested"),
        ("NFR-009", "Secrets shall be supplied through environment variables; .env files are gitignored and never committed; a launch-readiness script scans for hardcoded localhost URLs and placeholder secrets before deploy.", "Security", "Configured/documented"),
        ("NFR-010", "Password recovery shall require a single-use, expiring, hashed verification token before a password can be changed.", "Security", "Implemented - resolved since v1.0"),
        ("NFR-011", "An account shall be temporarily locked after repeated failed login attempts, independent of the IP-based rate limiter.", "Security", "Implemented - tested"),
        ("NFR-012", "Request/error logs shall never include headers, cookies or request bodies, which may carry bearer tokens, refresh cookies or user secrets - only route, method and user id are logged for context.", "Security", "Implemented"),
    ], "app middleware, auth/storage configuration, validators, error handler, rate-limit/requireAuth/requireRole/upload/password/passwordReset tests, and deployment documentation.")
    add_resolved_callout(doc, "NFR-010 resolved since v1.0", "v1.0 listed this as Not implemented. It is implemented as of 16 July 2026 - see Section 5.")

    # ---------------- 28. NFR Performance ----------------
    page_break(doc)
    add_page_title(doc, "28", "Non-Functional Requirements - Performance and Reliability")
    add_requirements(doc, [
        ("NFR-013", "List endpoints shall use server-side pagination and database filtering/sorting for scalable result sets.", "Performance", "Implemented"),
        ("NFR-014", "Aggregate dashboard queries shall avoid per-row N+1 query patterns; this is covered by a dedicated automated test.", "Performance", "Implemented - tested"),
        ("NFR-015", "The frontend shall lazy-load major role pages and produce optimized, code-split static assets through the Vite production build.", "Performance", "Implemented; build re-verified 20 Jul 2026"),
        ("NFR-016", "The API shall use response compression and support a pooled/direct PostgreSQL connection split (DATABASE_URL / DIRECT_URL) so short-lived hosted instances do not exhaust the connection limit of a serverless-friendly provider such as Neon.", "Performance", "Implemented/configured"),
        ("NFR-017", "The service shall expose a dependency-free liveness endpoint (GET /health) and a readiness endpoint (GET /api/health) that checks database and storage connectivity.", "Reliability", "Implemented - tested"),
        ("NFR-018", "Readiness failures shall return HTTP 503 with a per-component ok/error status object and no sensitive detail.", "Reliability", "Implemented - tested"),
        ("NFR-019", "The backend shall perform graceful shutdown on SIGTERM/SIGINT: stop accepting new connections, disconnect Prisma, and force-exit after a 10-second grace period if connections have not drained.", "Reliability", "Implemented"),
        ("NFR-020", "Every request shall be bounded by a request timeout (30 seconds) and a headers timeout (35 seconds, deliberately greater than the request timeout per Node's documented requirement).", "Reliability", "Implemented"),
        ("NFR-021", "A failed database write after a file upload shall remove the newly stored file to avoid an orphaned upload.", "Reliability", "Implemented - tested"),
        ("NFR-022", "Production uploads shall use persistent object storage (STORAGE_PROVIDER=s3); local filesystem storage is suitable only where the host volume is guaranteed persistent, which Railway/Render's default filesystem is not.", "Reliability", "Configurable; deployment caveat"),
        ("NFR-023", "A live deployment shall be verified for database migration, restart persistence, storage persistence and cross-domain session behavior before being treated as production-ready.", "Operations", "Local dev DB verified; hosted deployment not yet evidenced"),
    ], "pagination helpers/services, n-plus-one.test.ts, health.test.ts, Vite build output, index.ts shutdown handlers, storage providers, DEPLOYMENT.md.")

    # ---------------- 29. NFR Observability ----------------
    page_break(doc)
    add_page_title(doc, "29", "Non-Functional Requirements - Observability and Monitoring", "New section - the underlying hook points did not exist in v1.0.")
    add_requirements(doc, [
        ("NFR-024", "The backend shall emit structured, single-line JSON log entries (event, level, userId, time, env) so hosted log views or a jq pipeline can filter without parsing free text.", "Observability", "Implemented"),
        ("NFR-025", "The backend shall integrate Sentry (@sentry/node), initialized once at startup only when SENTRY_DSN is configured, and shall capture unhandled promise rejections, uncaught exceptions, and every error reaching the central Express error handler.", "Observability", "Implemented - inert without a DSN"),
        ("NFR-026", "The frontend shall integrate Sentry (@sentry/react), initialized once at startup only when VITE_SENTRY_DSN is configured, and shall capture every uncaught render error reaching the top-level ErrorBoundary.", "Observability", "Implemented - inert without a DSN"),
        ("NFR-027", "Neither monitoring integration shall transmit any data in an environment where its DSN environment variable is unset - both reportError() functions are documented no-ops in that case.", "Privacy", "Implemented - by design"),
        ("NFR-028", "An uncaught exception shall flush any pending Sentry event (bounded to 2 seconds) before the process exits, so the crash report that matters most is not silently lost.", "Observability", "Implemented"),
    ], "backend/src/utils/monitoring.ts, backend/src/index.ts, frontend/src/utils/monitoring.ts, frontend/src/main.tsx, backend/src/middleware/errorHandler.ts, frontend/src/components/ErrorBoundary.tsx.")
    add_gap_callout(doc, "Prepared, not connected", "No environment reviewed for this document has a SENTRY_DSN or VITE_SENTRY_DSN configured. The integration is wired correctly end-to-end but has never actually sent an event; connecting it requires only creating a Sentry account and setting the two environment variables - see DEPLOYMENT.md's monitoring checklist.")

    # ---------------- 30. NFR Usability ----------------
    page_break(doc)
    add_page_title(doc, "30", "Non-Functional Requirements - Usability and Accessibility")
    add_requirements(doc, [
        ("NFR-029", "The user interface shall use consistent role navigation, headings, colors, reusable controls and responsive layouts.", "Usability", "Implemented"),
        ("NFR-030", "Forms shall provide labels, validation feedback, saving state and cancellation/confirmation behavior.", "Usability", "Implemented"),
        ("NFR-031", "Data views shall provide loading skeletons, empty states, searchable/filterable lists and visible status indicators.", "Usability", "Implemented"),
        ("NFR-032", "Interactive dialogs and controls shall support keyboard escape/click-outside patterns and accessible labels where implemented.", "Accessibility", "Partially implemented"),
        ("NFR-033", "The application shall recover from uncaught frontend rendering failures with an error boundary and a reload action, without losing the user's authenticated session.", "Usability", "Implemented"),
        ("NFR-034", "A formal WCAG audit (keyboard traversal, focus order, color contrast, screen-reader labeling) should be performed before any academic or external submission.", "Accessibility", "Not performed"),
    ], "frontend components/hooks/styles, TypeScript structure, ErrorBoundary.tsx.")

    # ---------------- 31. NFR Maintainability ----------------
    page_break(doc)
    add_page_title(doc, "31", "Non-Functional Requirements - Maintainability and Quality")
    add_requirements(doc, [
        ("NFR-035", "The codebase shall use shared TypeScript types, zod validators, services, components and Zustand stores to reduce duplication.", "Maintainability", "Implemented"),
        ("NFR-036", "The backend shall expose OpenAPI/Swagger documentation, generated from route-level JSDoc comments, outside production.", "Maintainability", "Implemented"),
        ("NFR-037", "Database changes shall be versioned through Prisma migrations (10 migrations as of the evidence date) and validated (prisma validate, prisma format --check) before deployment.", "Maintainability", "Implemented"),
        ("NFR-038", "ESLint (typescript-eslint) shall pass with zero errors on both the frontend and backend.", "Quality", "Implemented - verified"),
        ("NFR-039", "The backend automated suite shall cover authentication, authorization, validation, pagination, uploads, health, audit coverage, N+1 avoidance, the Training Plan automation transaction, session/assignment feedback, announcements, password reset and token refresh - run against a mocked Prisma client with no live database required.", "Quality", "27 files / 140 tests - all passing, re-run 20 Jul 2026"),
        ("NFR-040", "The frontend automated suite shall cover demo-data fixture integrity, Demo Mode API-interception parity with the real backend, and the shared working-day/session-time calculation.", "Quality", "3 files / 29 tests - all passing, re-run 20 Jul 2026"),
        ("NFR-041", "An end-to-end Playwright suite shall exercise Demo Mode across all three roles.", "Quality", "1 spec file, demo-flows.spec.ts"),
        ("NFR-042", "A launch-readiness script (npm run launch:check) shall automate pre-deploy checks: env completeness, placeholder-secret detection, Prisma validation, migration safety review, build success for both apps, and the full backend test suite - in one command with no secrets printed.", "Operations", "Implemented"),
    ], "frontend/backend eslint config, Swagger config, Prisma migrations, backend/src/__tests__/*, frontend/src/__tests__/*, frontend/e2e/demo-flows.spec.ts, launch-check tooling.")

    # ---------------- 32. Data, Integrations ----------------
    page_break(doc)
    add_page_title(doc, "32", "Data, Integrations and Constraints")
    doc.add_paragraph("Core record groups", style="Heading 2")
    add_definition_rows(doc, [
        ("Identity", "Role, Permission, RolePermission, User, UserProfile, UserInvite, RefreshToken, PasswordResetToken"),
        ("Curriculum templates", "TrainingPlan, TrainingPlanSession, TrainingPlanAssignment, TrainingPlanResource, TrainingPlanAnnouncement"),
        ("Training operations", "Batch, BatchTrainee, Assignment, AssignmentBatch, Submission, SubmissionAttachment"),
        ("Delivery", "Session, Attendance, Resource"),
        ("Feedback (three distinct kinds)", "FeedbackEntry (performance ratings); SessionFeedbackForm/Submission and AssignmentFeedbackForm/Submission (external form links)"),
        ("Engagement", "Announcement, AnnouncementRead, DiscussionThread, DiscussionMessage (dormant)"),
        ("Operations", "AuditLog, NotificationRead"),
    ])
    doc.add_paragraph("Implemented integration boundaries", style="Heading 2")
    add_bullets(doc, [
        "PostgreSQL 17/Prisma for durable business data. A local instance is connected, migrated (10 migrations applied) and seeded, and has been exercised end-to-end (login/JWT, announcements CRUD with real per-user read tracking). A hosted/production database (the documented Neon path) has not yet been proven.",
        "Local disk or S3 storage through a provider interface; Cloudinary is not implemented.",
        "Console-log email provider only. No real SES/Postmark/Resend connection exists - invite and password-reset emails are not actually delivered outside a manual/dev workaround.",
        "Browser mailto: links for selected contact actions; no internal messaging service exists.",
        "No WebSocket, push notification or background job/scheduler service exists anywhere in the codebase.",
        "Sentry error tracking is wired at both layers but inert until a DSN is supplied (Section 29).",
        "Deployment files target Vercel plus Railway/Render and a Neon-compatible database, and a full deployment guide with a pre-deploy readiness script exists, but the repository does not yet prove a completed live deployment.",
    ], compact=True)
    add_gap_callout(doc, "Dormant schema", "Discussion tables exist in Prisma (Section 22), but no backend routes/services exist for them. Their presence in the schema must not be interpreted as completed functionality.")

    # ---------------- 33. Known Gaps ----------------
    page_break(doc)
    add_page_title(doc, "33", "Known Gaps and Future Enhancements")
    doc.add_paragraph("Resolved since v1.0 (13 July 2026)", style="Heading 2")
    add_bullets(doc, [
        "Forgot Password now issues and verifies a single-use, hashed, one-hour token instead of accepting a bare email plus new password (Section 5).",
        "Announcements are now database-backed with real per-user read tracking and full Demo Mode parity (Section 17) - the module v1.0 called the last remaining mock-data holdout.",
        "The pageSize:200 vs backend cap:100 mismatch that 400'd every real-mode list call using the fetch-everything pattern has been fixed (cap raised to 500 backend and frontend together).",
        "A migration file that had silently been saved as UTF-16 (unparseable by Prisma's migration engine) has been identified and rewritten as UTF-8.",
        "Assignment-level feedback links, flagged as an unresolved conflict in the 15 July 2026 audit, were implemented the next day (Section 15).",
        "A local PostgreSQL instance is now connected and has been exercised end-to-end for the first time (login/JWT persistence, real announcements round-trip).",
    ], compact=True)
    doc.add_paragraph("Confirmed gaps still open", style="Heading 2")
    add_bullets(doc, [
        "No real email provider sends invitations, password resets, or any other notification - the console-log provider is the only one implemented.",
        "No hosted/production PostgreSQL has been proven; the local dev connection is not equivalent to a live Neon deployment.",
        "S3 file storage is implemented against the provider interface but not exercised against a real bucket in this review; Railway/Render's default filesystem is not persistent, so local storage would lose files on redeploy in that environment.",
        "Scheduled reports remain configuration-only; no backend scheduler or job runner executes them.",
        "No WebSocket/push layer exists despite the Real-time Analytics label in the Admin navigation.",
        "Discussions has a dormant schema with no backend routes and no UI; it is intentionally out of scope, not broken.",
        "Sentry/error tracking is wired end-to-end at both layers but has never sent an event in any reviewed environment - no DSN is configured anywhere.",
        "Formal WCAG accessibility testing has not been performed.",
        "No avatar upload endpoint exists, though the data model reserves fields for one.",
        "The Role/Permission/RolePermission tables and requirePermission() middleware exist but are not exposed through any admin UI, and most routes still authorize by role name rather than the granular permission set.",
        "The submission-roster and resource-visibility least-privilege caveats carried forward from the 15 July 2026 audit were not independently re-verified for this revision.",
        "No completed production deploy, live migration run, backup/restore drill, or connected uptime monitor has been evidenced.",
    ], compact=True)
    doc.add_paragraph("Prioritized future enhancements", style="Heading 2")
    add_definition_rows(doc, [
        ("P1 - Communication", "Connect a real email provider (SES/Postmark/Resend) so invite and password-reset tokens actually reach users."),
        ("P1 - Operations", "Execute the documented Neon + Railway/Render + Vercel deployment and run the full manual launch checklist against it."),
        ("P1 - Security review", "Independently re-verify the submission-roster and resource-visibility least-privilege caveats; complete a formal security review."),
        ("P2 - Automation", "Add a background job runner for scheduled report generation and delivery, with run history."),
        ("P2 - Observability", "Provision a Sentry DSN in a real deployment and confirm a test event arrives; connect an uptime monitor to GET /api/health."),
        ("P2 - Storage", "Provision and exercise a real S3 bucket end to end (upload, download, redeploy-survival) before relying on it for graded/production submissions."),
        ("P3 - Access model", "Either build an admin UI over the existing Role/Permission tables, or retire them if the role-name model is the long-term design."),
        ("P3 - Experience", "Add push/live updates, formal WCAG testing, an avatar upload endpoint, and richer audit export controls."),
        ("P3 - Messaging", "Implement the dormant discussion schema as a real feature, or remove it from the schema if it will not be built."),
    ])
    add_screenshot_placeholder(doc, "Future evidence", "replace with final production deployment, monitoring dashboard, and accessibility-audit evidence once available")

    # ---------------- 34. Conclusion ----------------
    page_break(doc)
    add_page_title(doc, "34", "Conclusion and Verification Summary")
    add_text(doc, "The implemented Trainee Portal now provides a substantially more complete foundation than it did at v1.0. The core business workflow - a fixed Training Plan template copied into an independent, editable set of records for every new batch - is fully implemented and exercised by both automated tests and manual verification against a real database. Around it, the portal covers authenticated role-based access, curriculum templating, batch and assignment operations, trainee submissions, grading, sessions, attendance, resources, three distinct feedback mechanisms, database-backed announcements, audit activity, notifications, dashboards and reporting views.")
    add_text(doc, "The strongest implementation areas are the typed API surface (89 REST endpoints), the relational data model (30+ Prisma models), the automated test suite (169 passing tests across 30 files), the training-plan-to-batch instantiation transaction, and the security posture around authentication (rate limiting, account lockout, hashed tokens, a now-fixed password-recovery flow). The remaining delivery risks are stated plainly rather than hidden: no real email provider is connected, no hosted production database has been proven, scheduled reports and real-time updates do not run, and error tracking is wired but has never actually reported an event.")
    doc.add_paragraph("Verification performed for this revision", style="Heading 2")
    add_definition_rows(doc, [
        ("Repository analysis", "Full Prisma schema, all 16 backend route files (89 endpoints), all 16 validator files, all middleware, config defaults, frontend page/store structure, and the DEPLOYMENT.md/USER_FLOWS.md/AUDIT_REPORT.md/CLAUDE_PROJECT_CONTEXT.md project documents were read directly for this revision."),
        ("Production build", "Re-run and passed for both frontend (tsc -b && vite build) and backend (tsc) on 20 July 2026."),
        ("Automated tests", "Re-run and passed on 20 July 2026: backend 27 files / 140 tests; frontend 3 files / 29 tests. Figures were not carried forward from any prior note."),
        ("End-to-end", "1 Playwright spec (Demo Mode, all three roles) is committed to the repository; a separate one-off manual/browser verification pass (58 PASS / 4 PARTIAL / 0 FAIL against a detailed checklist) was performed on 15 July 2026 and is documented in AUDIT_REPORT.md."),
        ("Live environment", "A local PostgreSQL instance has been connected, migrated and seeded, and exercised end-to-end for login/JWT and real Announcements CRUD. A hosted/production environment (Neon/Railway-Render/Vercel) has not been proven - no claim is made that it is operational."),
        ("Document scope", "141 functional requirements, 40 business rules and 42 non-functional requirements documented from the as-is implementation, plus a full data dictionary and API endpoint appendix that did not exist at v1.0."),
    ])
    doc.add_paragraph("Recommended sign-off statement", style="Heading 2")
    p = doc.add_paragraph()
    border_paragraph(p, color="9FB3C8", fill=LIGHT_BLUE)
    r = p.add_run("This FRD accurately represents the reviewed implementation as of 20 July 2026, subject to the explicitly listed gaps in Section 33 and the absence of live production deployment verification.")
    set_run_font(r, size=9.5, italic=True, color=NAVY)
    doc.add_paragraph("Final screenshot checklist", style="Heading 2")
    add_bullets(doc, [
        "Use a consistent browser size and zoom; remove personal information before capturing.",
        "Capture login, each role dashboard, Training Plan detail, batch onboarding, assignment/submission, calendar, resource, both feedback-form features, performance feedback, announcements, audit, and report screens.",
        "Replace every placeholder in this document only after final data and styling are stable.",
    ], compact=True)

    # ---------------- Appendix A: Data Dictionary ----------------
    page_break(doc)
    add_page_title(doc, "A", "Appendix A - Data Dictionary", "Every Prisma model in the current schema (backend/prisma/schema.prisma), grouped to match Section 32.")
    doc.add_paragraph("A.1 Identity and access", style="Heading 2")
    add_model_table(doc, "Role / Permission / RolePermission", "The granular permission model - present in the schema, not yet exposed through any admin UI (Section 33).", [
        ("Role.name", "admin / facilitator / trainee - the three role names route authorization actually checks."),
        ("Permission.key", "A granular capability key; joined to roles via RolePermission."),
    ])
    add_model_table(doc, "User", "Every account: administrators, facilitators and trainees share one table, distinguished by roleId.", [
        ("email, passwordHash", "Unique login identity; bcrypt hash, never returned by the API."),
        ("isActive, deletedAt", "Soft-delete pattern - a deleted or inactive user cannot authenticate."),
        ("failedLoginAttempts, lockedUntil", "Back the account-lockout rule (BR-002)."),
        ("lastLoginAt", "Set on every successful login."),
    ])
    add_model_table(doc, "UserProfile", "One-to-one extension of User with optional contact/company detail and reserved avatar fields.", [
        ("phone, location, company, department, idNumber", "All optional; shown read-only in Account Settings by role."),
        ("avatarStorageKey / avatarMimeType / avatarSizeBytes / avatarUpdatedAt", "Reserved for a future avatar-upload feature; no route sets them today (FR-022)."),
    ])
    add_model_table(doc, "UserInvite", "An administrator-issued, single-use invitation.", [
        ("tokenHash", "SHA-256 hash of the raw invite token; the raw value is never stored."),
        ("status", "Pending / Accepted / Expired / Revoked."),
        ("expiresAt", "7 days from creation (BR-006)."),
    ])
    add_model_table(doc, "RefreshToken", "One row per active or historical refresh-token session.", [
        ("tokenHash", "SHA-256 hash; rememberMe controls whether the 7-day or 30-day TTL applies."),
        ("revokedAt", "Set on logout or when the owning user is soft-deleted."),
    ])
    add_model_table(doc, "PasswordResetToken", "Backs the token-verified forgot-password flow (Section 5).", [
        ("tokenHash, expiresAt, usedAt", "Single-use, 1-hour TTL, hashed at rest - identical pattern to invites."),
    ])

    page_break(doc)
    doc.add_paragraph("A.2 Curriculum templates (Training Plans)", style="Heading 2")
    add_model_table(doc, "TrainingPlan", "The org's fixed curriculum template - exactly two rows exist (BA BTech, BA MBA).", [
        ("code, name, description, durationMonths", "Plan identity and length (default 2 months)."),
        ("defaultSessionStartMinute / EndMinute", "Minutes from midnight; default 870/990 = 14:30-16:30."),
        ("defaultAssignmentStartMinute / DeadlineMinute", "Default 570/1439 = 09:30 open / 23:59 due."),
    ])
    add_model_table(doc, "TrainingPlanSession", "One template session row per scheduled class in the plan.", [
        ("dayOffset", "Working-day index from the batch's start date (BR-016)."),
        ("startMinute, endMinute", "Must satisfy endMinute > startMinute (BR-017, validated)."),
        ("feedbackFormUrl", "Optional default, copied onto every instantiated Session's feedback form."),
    ])
    add_model_table(doc, "TrainingPlanAssignment", "One template assignment row, optionally tied to a template session.", [
        ("dueDayOffset", "Working-day index the generated assignment's deadline is computed from."),
        ("relatedSessionId", "Optional link to a TrainingPlanSession, carried onto the instantiated Assignment."),
    ])
    add_model_table(doc, "TrainingPlanResource / TrainingPlanAnnouncement", "Template-level defaults copied onto every batch generated from the plan.", [
        ("TrainingPlanResource.url", "External link; instantiated Resource rows created from this always populate externalUrl, never storageKey."),
        ("TrainingPlanAnnouncement.priority", "Normal / Important / Critical, copied onto the instantiated Announcement."),
    ])

    page_break(doc)
    doc.add_paragraph("A.3 Training operations", style="Heading 2")
    add_model_table(doc, "Batch", "One cohort - the central record everything else scopes against.", [
        ("code", "Unique; program/track are enum-derived from the linked TrainingPlan (BR-015)."),
        ("trainingPlanId", "Immutable after creation (BR-021)."),
        ("facilitatorId", "Optional Trainer (BR-023)."),
        ("startMonth, endDate", "endDate is always server-computed (BR-020)."),
        ("archivedAt, deletedAt", "Soft-archive and soft-delete are tracked separately."),
    ])
    add_model_table(doc, "BatchTrainee", "Enrollment join table, keyed by (batchId, traineeId).", [
        ("enrolledAt, removedAt", "removedAt is set on removal rather than deleting the row (BR-024)."),
    ])
    add_model_table(doc, "Assignment / AssignmentBatch", "A piece of work, optionally assigned to multiple batches at once.", [
        ("Assignment.batchId", "Primary batch, retained for backward compatibility (BR-025)."),
        ("AssignmentBatch", "Full many-to-many set; every assignment has at least one row here."),
        ("trainingPlanAssignmentId, sessionId", "Nullable back-references to a template origin and/or a related Session."),
        ("attachmentStorageKey (unique)", "The optional instructions file."),
    ])
    add_model_table(doc, "Submission / SubmissionAttachment", "A trainee's work against one assignment, with full attachment history.", [
        ("Submission unique(assignmentId, traineeId)", "Enforces one submission per trainee per assignment (BR-026)."),
        ("grade", "Decimal(5,2), 0-100 (BR-029)."),
        ("SubmissionAttachment.isCurrent", "False for superseded files, retained for history (BR-028)."),
    ])

    page_break(doc)
    doc.add_paragraph("A.4 Delivery, feedback, engagement and operations", style="Heading 2")
    add_model_table(doc, "Session / Attendance", "A scheduled class and its per-trainee attendance records.", [
        ("Session.durationMinutes", "Default 120 (2:30-4:30 PM standard block)."),
        ("Session.trainingPlanSessionId", "Nullable back-reference to the template session it was generated from."),
        ("Attendance unique(sessionId, traineeId)", "One record per trainee per session (BR-030)."),
    ])
    add_model_table(doc, "Resource", "A learning resource, either an uploaded file or an external link.", [
        ("storageKey (unique) vs externalUrl", "Exactly one is populated (BR-031)."),
        ("verified, downloadCount", "verified gates trainee visibility in the UI; downloadCount increments on every download."),
        ("trainingPlanResourceId", "Nullable back-reference to a template resource."),
    ])
    add_model_table(doc, "SessionFeedbackForm / SessionFeedbackSubmission", "External form link attached to a Session, plus per-user completion tracking.", [
        ("formUrl, audience", "URL-validated link; audience is Trainees/Facilitators/Both (BR-036)."),
        ("SessionFeedbackSubmission unique(formId, submitterId)", "Idempotent per-user completion marker (BR-038)."),
    ])
    add_model_table(doc, "AssignmentFeedbackForm / AssignmentFeedbackSubmission", "The same pattern as above, attached directly to an Assignment instead of a Session.", [
        ("assignmentId (unique)", "One feedback form per assignment, independent of any related session's form (FR-099)."),
    ])
    add_model_table(doc, "FeedbackEntry", "Rating-based performance feedback between a facilitator and a trainee.", [
        ("direction", "FacilitatorToTrainee or TraineeToFacilitator, stamped automatically (BR-035 area)."),
        ("rating", "Integer 1-5; no update/delete route exists - append-only."),
    ])
    add_model_table(doc, "Announcement / AnnouncementRead", "A posted notice, global or batch-scoped, with per-user read tracking.", [
        ("batchId (nullable)", "Null means global; a facilitator may only post batch-scoped announcements (BR-039)."),
        ("trainingPlanAnnouncementId", "Nullable back-reference to a template default announcement."),
        ("AnnouncementRead", "Composite key (announcementId, userId); independent of the announcement's own edit history (BR-040)."),
    ])
    add_model_table(doc, "DiscussionThread / DiscussionMessage", "Dormant - present in the schema, no backend routes, no UI (Section 22).", [
        ("authorRoleSnapshot", "Captures the author's role at message time even if their role later changes; unused by any route today."),
    ])
    add_model_table(doc, "AuditLog / NotificationRead", "The append-only activity log and its per-user notification read state.", [
        ("eventType, module, previousValue, newValue", "Structured enough to drive both the Admin audit log screen and the derived notification feed."),
        ("NotificationRead", "Composite key (auditLogId, userId); never mutates the audit entry itself (BR-034)."),
    ])

    # ---------------- Appendix B: API Endpoint Reference ----------------
    page_break(doc)
    add_page_title(doc, "B", "Appendix B - API Endpoint Reference", "All 89 REST endpoints under /api, grouped by module and read directly from the route source files' OpenAPI annotations.")
    doc.add_paragraph("B.1 Authentication, users and health", style="Heading 2")
    add_endpoint_table(doc, [
        ("POST /auth/login", "Public, rate-limited", "Authenticate with email/password; returns access token + user profile."),
        ("POST /auth/logout", "Authenticated", "Revoke the current refresh token and clear the session cookie."),
        ("POST /auth/refresh", "Refresh cookie", "Exchange a valid refresh token cookie for a new access token."),
        ("GET /auth/me", "Authenticated", "Get the currently authenticated user."),
        ("POST /auth/invite", "Admin", "Invite a new user by email."),
        ("POST /auth/invite/accept", "Public + invite token", "Accept an invite and set a password."),
        ("POST /auth/forgot-password", "Public, rate-limited", "Request a password-reset email; always succeeds (no enumeration)."),
        ("POST /auth/reset-password", "Public + reset token, rate-limited", "Set a new password using an emailed single-use reset token."),
        ("POST /auth/change-password", "Authenticated", "Change the current user's password (requires current password)."),
        ("GET /users", "Admin", "List users - paginated, filterable, searchable, sortable."),
        ("GET /users/me", "Authenticated", "Get the current user's profile."),
        ("PATCH /users/me", "Authenticated", "Update the current user's own name/email/phone/location."),
        ("GET /users/:id", "Admin or self", "Get a user by id."),
        ("PATCH /users/:id", "Admin", "Update any user, including role/isActive."),
        ("DELETE /users/:id", "Admin", "Soft-delete a user."),
        ("GET /api/health", "Public", "Readiness check - server, database and file storage connectivity."),
        ("GET /health", "Public", "Dependency-free liveness probe (top-level, not under /api)."),
    ])

    page_break(doc)
    doc.add_paragraph("B.2 Training Plans and Batches", style="Heading 2")
    add_endpoint_table(doc, [
        ("GET /training-plans", "Authenticated", "List all Training Plans (currently just BA BTech and BA MBA)."),
        ("GET /training-plans/:id", "Authenticated", "Get a plan's full template (sessions, assignments, resources, announcements)."),
        ("PATCH /training-plans/:id", "Admin", "Edit a plan's name/description/duration/default timing."),
        ("POST /training-plans/:id/sessions", "Admin", "Add a session to a plan's template."),
        ("PATCH /training-plans/:id/sessions/:sessionId", "Admin", "Edit a template session."),
        ("DELETE /training-plans/:id/sessions/:sessionId", "Admin", "Remove a template session."),
        ("POST /training-plans/:id/assignments", "Admin", "Add an assignment to a plan's template."),
        ("PATCH /training-plans/:id/assignments/:assignmentId", "Admin", "Edit a template assignment."),
        ("DELETE /training-plans/:id/assignments/:assignmentId", "Admin", "Remove a template assignment."),
        ("POST /training-plans/:id/resources", "Admin", "Add a resource to a plan's template."),
        ("PATCH /training-plans/:id/resources/:resourceId", "Admin", "Edit a template resource."),
        ("DELETE /training-plans/:id/resources/:resourceId", "Admin", "Remove a template resource."),
        ("POST /training-plans/:id/announcements", "Admin", "Add a default announcement to a plan's template."),
        ("PATCH /training-plans/:id/announcements/:announcementId", "Admin", "Edit a template announcement."),
        ("DELETE /training-plans/:id/announcements/:announcementId", "Admin", "Remove a template announcement."),
        ("GET /batches", "Authenticated", "List batches - paginated, filterable, searchable, sortable."),
        ("POST /batches", "Admin", "Create a batch from a Training Plan."),
        ("GET /batches/:id", "Authenticated", "Get a batch by id."),
        ("PATCH /batches/:id", "Admin", "Update a batch."),
        ("DELETE /batches/:id", "Admin", "Archive/soft-delete a batch."),
        ("GET /batches/:id/metrics", "Authenticated", "Derived performance metrics for a batch."),
        ("GET /batches/:id/trainee-stats", "Admin or owning Facilitator", "Per-trainee stats for a batch."),
        ("GET /batches/:id/trainees", "Authenticated", "List trainees enrolled in a batch."),
        ("POST /batches/:id/trainees", "Admin, Facilitator", "Enroll a trainee in a batch."),
        ("DELETE /batches/:id/trainees/:traineeId", "Admin, Facilitator", "Remove a trainee from a batch."),
    ])

    page_break(doc)
    doc.add_paragraph("B.3 Assignments, submissions, sessions and attendance", style="Heading 2")
    add_endpoint_table(doc, [
        ("GET /assignments", "Authenticated", "List assignments - paginated, filterable by batch/status."),
        ("POST /assignments", "Admin, Facilitator", "Create an assignment for one or more batches, with an optional file."),
        ("GET /assignments/:id", "Authenticated", "Get an assignment by id."),
        ("PATCH /assignments/:id", "Admin, owner Facilitator", "Update an assignment, its batches, or its file."),
        ("DELETE /assignments/:id", "Admin, owner Facilitator", "Soft-delete an assignment."),
        ("GET /assignments/:id/attachment", "Scoped users", "View/download an assignment's instructions file."),
        ("GET /assignments/:id/submissions", "Authenticated", "List submissions, incl. a Not-Started placeholder per trainee."),
        ("POST /assignments/:id/submissions", "Trainee", "Submit or resubmit the current trainee's own work."),
        ("GET /submissions/:id", "Owning trainee, facilitator or admin", "Get a submission."),
        ("PATCH /submissions/:id", "Admin, owning Facilitator", "Grade a submission."),
        ("POST /submissions/:id/attachments", "Trainee (own submission)", "Upload a file attachment to a submission."),
        ("GET /submissions/:id/attachments/:attachmentId", "Scoped users", "Download a submission attachment."),
        ("GET /sessions", "Authenticated", "List sessions - paginated, filterable by batch/status."),
        ("POST /sessions", "Admin, Facilitator", "Create a session."),
        ("GET /sessions/:id", "Authenticated", "Get a session by id."),
        ("PATCH /sessions/:id", "Admin, owning Facilitator", "Update a session."),
        ("DELETE /sessions/:id", "Admin, owning Facilitator", "Soft-delete a session."),
        ("GET /sessions/:id/attendance", "Admin, owning Facilitator", "List attendance records for a session."),
        ("PUT /sessions/:id/attendance", "Admin, owning Facilitator", "Bulk mark attendance for a session."),
        ("PATCH /attendance/:id", "Admin, owning Facilitator", "Correct a single attendance record."),
        ("GET /calendar", "Authenticated", "Normalized sessions + assignment-deadline feed, scoped by role."),
    ])

    page_break(doc)
    doc.add_paragraph("B.4 Resources, feedback (all three kinds), announcements and notifications", style="Heading 2")
    add_endpoint_table(doc, [
        ("GET /resources", "Authenticated", "List resources - paginated, filterable by batch/category/verified."),
        ("POST /resources", "Admin, Facilitator", "Upload a learning resource."),
        ("GET /resources/:id", "Authenticated", "Get a resource's metadata."),
        ("PATCH /resources/:id", "Admin, owner uploader", "Update resource metadata."),
        ("DELETE /resources/:id", "Admin, owner uploader", "Soft-delete a resource."),
        ("GET /resources/:id/download", "Authenticated", "Download a resource file; increments its download count."),
        ("GET /sessions/:id/feedback-form", "Admin, owning Facilitator, enrolled Trainee", "Get a session's feedback-form link and stats."),
        ("POST /sessions/:id/feedback-form", "Admin, owning Facilitator", "Attach a feedback-form link to a session."),
        ("PATCH /sessions/:id/feedback-form", "Admin, owning Facilitator", "Edit a session's feedback form."),
        ("DELETE /sessions/:id/feedback-form", "Admin, owning Facilitator", "Remove a session's feedback form."),
        ("POST /sessions/:id/feedback-form/submit", "Eligible Trainee or Facilitator", "Mark session feedback as submitted (idempotent)."),
        ("GET /assignments/:id/feedback-form", "Admin, owning Facilitator, enrolled Trainee", "Get an assignment's feedback-form link and stats."),
        ("POST /assignments/:id/feedback-form", "Admin, owning Facilitator", "Attach a feedback-form link to an assignment."),
        ("PATCH /assignments/:id/feedback-form", "Admin, owning Facilitator", "Edit an assignment's feedback form."),
        ("DELETE /assignments/:id/feedback-form", "Admin, owning Facilitator", "Remove an assignment's feedback form."),
        ("POST /assignments/:id/feedback-form/submit", "Eligible Trainee or Facilitator", "Mark assignment feedback as submitted (idempotent)."),
        ("GET /feedback", "Authenticated (scoped)", "List performance-feedback entries."),
        ("POST /feedback", "Admin, Facilitator, Trainee (scoped)", "Record performance feedback; append-only."),
        ("GET /feedback/:id", "Authenticated (scoped)", "Get a feedback entry by id."),
        ("GET /announcements", "Authenticated", "List announcements visible to the caller."),
        ("POST /announcements", "Admin, Facilitator", "Publish an announcement."),
        ("PATCH /announcements/:id", "Author or Admin", "Update an announcement."),
        ("DELETE /announcements/:id", "Author or Admin", "Soft-delete an announcement."),
        ("POST /announcements/:id/read", "Authenticated", "Mark an announcement as read for the current user."),
        ("GET /notifications", "Authenticated", "List notifications derived from the audit log."),
        ("POST /notifications/:id/read", "Authenticated", "Mark a single notification as read."),
        ("POST /notifications/read-all", "Authenticated", "Mark all notifications as read."),
    ])

    # ---------------- Appendix C: Environment Configuration ----------------
    page_break(doc)
    add_page_title(doc, "C", "Appendix C - Environment Configuration Reference", "Selected runtime configuration read from backend/src/config/env.ts; every value shown is the code default unless overridden by an environment variable.")
    add_definition_rows(doc, [
        ("JWT_ACCESS_EXPIRES_IN", "15m - access token lifetime."),
        ("REFRESH_TOKEN_TTL_DAYS / REMEMBER_TTL_DAYS", "7 days default / 30 days with Remember Me."),
        ("BCRYPT_SALT_ROUNDS", "12 (configurable 8-15)."),
        ("LOGIN_MAX_ATTEMPTS / LOGIN_LOCKOUT_MINUTES", "5 attempts / 15-minute lockout."),
        ("MAX_UPLOAD_SIZE_MB", "10 MB per file."),
        ("STORAGE_PROVIDER", "local (default) or s3; s3 requires AWS_REGION and AWS_S3_BUCKET."),
        ("SENTRY_DSN / VITE_SENTRY_DSN", "Unset by default in every reviewed environment; monitoring is inert until set (Section 29)."),
        ("EXPOSE_AUTH_TOKENS", "Defaults to true outside production, false in production - controls whether invite/reset tokens are echoed in API responses (BR-009)."),
        ("CORS_ORIGIN", "Comma-separated allowlist of origins permitted to call the API with credentials."),
        ("API rate limit (general)", "300 requests / 15 minutes per client."),
        ("Login rate limit", "10 attempts / 15 minutes per client, counting only failures."),
        ("Forgot-password rate limit", "5 requests / 60 minutes per client."),
    ])
    add_gap_callout(doc, "Not a secrets list", "This appendix lists variable names and code-level defaults only, exactly as documented in backend/.env.example. No real credential, connection string, or secret value from any environment is reproduced anywhere in this document.")

    # ---------------- Appendix D: Glossary ----------------
    page_break(doc)
    add_page_title(doc, "D", "Appendix D - Glossary")
    add_definition_rows(doc, [
        ("Training Plan", "The organization's fixed curriculum template (exactly two: BA BTech, BA MBA); see Section 7."),
        ("Batch", "One cohort of trainees running a Training Plan on a specific schedule, from a specific start date."),
        ("Instantiation", "The one-time transaction that copies a Training Plan's sessions, assignments, resources and default announcements into a new batch's own editable records."),
        ("Session Feedback / Assignment Feedback", "External form-link features (Section 14-15) that track click-through/submission state only - distinct from performance feedback."),
        ("Performance Feedback (FeedbackEntry)", "Rating-based feedback exchanged between a facilitator and a trainee (Section 16)."),
        ("Demo Mode", "A frontend-only interception layer that fakes every API call with fixture data, entered explicitly via the login page's View as buttons; never touches a real database."),
        ("Working day", "Monday through Friday; the schedule generator skips weekends and rolls a weekend start date forward to the next Monday."),
        ("Minute-of-day field", "An integer 0-1440 representing a wall-clock time (e.g. 870 = 14:30), used throughout the Training Plan template instead of a literal time string."),
        ("Soft delete", "Setting a deletedAt (or removedAt for enrollments) timestamp instead of removing the row, preserving history."),
        ("Readiness vs. liveness", "GET /health is a fast, dependency-free liveness probe; GET /api/health is a readiness probe that checks the database and storage provider and can return 503."),
    ])

    # Core metadata and deterministic document properties.
    props = doc.core_properties
    props.title = "Trainee Portal Functional Requirements Document v2.0"
    props.subject = "As-is functional requirements based on implemented source code"
    props.author = "Business Analyst"
    props.keywords = "FRD, Trainee Portal, functional requirements, business analysis, Training Plans"
    props.comments = "Generated from repository evidence; known gaps are explicitly identified. Supersedes v1.0."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
